"""
Simple README to PDF converter using Word automation
Converts: README_FIRST.md → README_FIRST.docx → README_FIRST.pdf
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client

def parse_markdown_to_docx(md_path, docx_path):
    """Convert markdown to Word document with basic formatting"""

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 headers (# Header)
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)  # Navy blue
            i += 1
            continue

        # H2 headers (## Header)
        if line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)  # Blue
            i += 1
            continue

        # H3 headers (### Header)
        if line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)  # Light blue
            i += 1
            continue

        # H4 headers (#### Header)
        if line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            i += 1
            continue

        # Horizontal rules (---)
        if line.startswith('---'):
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Code blocks (```)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue

        # Bullet lists (- item or * item)
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue

        # Numbered lists (1. item)
        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(line[line.index('.') + 2:], style='List Number')
            i += 1
            continue

        # Regular paragraphs
        # Handle bold (**text**) and links ([text](url))
        paragraph_text = line
        paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'\1', paragraph_text)  # Remove ** for bold
        paragraph_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', paragraph_text)  # Keep link text only

        p = doc.add_paragraph(paragraph_text)

        # Make text bold if it had **
        if '**' in line:
            for run in p.runs:
                if line.count('**') >= 2:
                    run.bold = True

        i += 1

    doc.save(docx_path)
    print(f"[OK] Created DOCX: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using Word automation"""

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        doc = word.Documents.Open(str(docx_path.absolute()))
        doc.SaveAs2(str(pdf_path.absolute()), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        print(f"[OK] Created PDF: {pdf_path}")
    finally:
        word.Quit()

def main():
    """Main conversion function"""
    base_dir = Path(__file__).parent

    md_path = base_dir / "README_FIRST.md"
    docx_path = base_dir / "README_FIRST.docx"
    pdf_path = base_dir / "README_FIRST.pdf"

    if not md_path.exists():
        print(f"[ERROR] {md_path} not found")
        sys.exit(1)

    print("Converting README_FIRST.md to PDF...")
    print(f"Step 1: Markdown -> DOCX")
    parse_markdown_to_docx(md_path, docx_path)

    print(f"Step 2: DOCX -> PDF")
    convert_docx_to_pdf(docx_path, pdf_path)

    print(f"\n[SUCCESS] Created {pdf_path}")

    # Optionally clean up intermediate DOCX
    # docx_path.unlink()

if __name__ == "__main__":
    try:
        main()
        print("\n[OK] README to PDF conversion SUCCESS")
        sys.exit(0)
    except Exception as e:
        print(f"\n[FAIL] README to PDF conversion FAILED: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
