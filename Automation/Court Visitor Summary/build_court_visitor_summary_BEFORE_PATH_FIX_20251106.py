#!/usr/bin/env python3
"""
Court Visitor Summary Form Generator

Top half: three columns (Ward | Guardian 1 | Guardian 2)
Bottom half: full-width ruled lines for notes
Excel opened read-only; optional --print and --open
"""
from __future__ import annotations
import argparse, os, re, sys
from datetime import datetime, date
from dateutil import parser as dtparse
from typing import List, Dict
import pandas as pd

# -------- deps --------
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError:
    print("ERROR: Missing dependency 'python-docx'. Install it with:\n  py -3 -m pip install python-docx")
    sys.exit(1)

try:
    from openpyxl import load_workbook
except ModuleNotFoundError:
    print("ERROR: Missing dependency 'openpyxl'. Install it with:\n  py -3 -m pip install openpyxl")
    sys.exit(1)

# Optional printing via Word
try:
    import win32com.client  # type: ignore
except Exception:
    win32com = None

# -------- config --------
# Use Excel file from app directory
WORKBOOK_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "App Data", "ward_guardian_info.xlsx")
SHEET_NAME = None
# Output to user-friendly location in app directory
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "Visit Summary Forms")

FONT_NAME = "Calibri"
FONT_SIZE_PT = 11
HEADER_FONT_SIZE_PT = 12   # small, just a tad larger than body
SUBHEAD_FONT_SIZE_PT = 11
MARGIN_IN = 0.4            # small margins
NOTES_ROWS = 26

def _clean(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

# Covers your sample headers & common typos
COLUMN_ALIASES = {
    "caseno": ["caseno", "causeno", "case", "case no", "cause no", "cause"],
    "visitdate": ["visitdate", "visit date", "date"],
    "visittime": ["visittime", "visit time", "time"],
    "wardlast": ["wardlast", "ward last", "wlast"],
    "wardfirst": ["wardfirst", "ward first", "wfirst"],
    "wtele": ["wtele", "wphone", "ward phone", "wardtelephone"],
    "liveswith": ["liveswith", "lives with"],
    "waddress": ["waddress", "ward address", "waddr"],
    "wdob": ["wdob", "ward dob", "warddob", "ward birth"],

    "guardian1": ["guardian1", "g1", "guardian"],
    "gaddress": ["gaddress", "guardian address", "g1address", "g addr"],
    "gemail": ["gemail", "g1email", "guardian email", "g e-mail"],
    "gtele": ["gtele", "g1tele", "guardian phone", "gphone"],
    "Relationship": ["Relationship", "relationship", "g1relationship", "g relationship"],
    "gdob": ["gdob", "g1dob", "guardian dob"],

    "Guardian2": ["Guardian2", "guardian2", "g2", "guardian 2"],
    "g2 address": ["g2 address", "g2address", "guardian2 address", "guardian 2 address"],
    "g2eamil": ["g2eamil", "g2email", "guardian2 email", "guardian 2 email"],
    "g2tele": ["g2tele", "guardian2 phone", "g2 phone"],
    "g2Relationship": ["g2Relationship", "g2relationship", "guardian2 relationship", "g2 relation"],
    "g2dob": ["g2dob", "guardian2 dob"],

    "miles": ["miles", "distance"],
}

# -------- helpers --------

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    colmap: Dict[str, str] = {}
    lowered = {str(c).lower().strip(): c for c in df.columns}
    for canon, variants in COLUMN_ALIASES.items():
        for v in variants:
            key = v.lower()
            if key in lowered:
                colmap[lowered[key]] = canon
                break
    return df.rename(columns=colmap)

def coalesce(row: pd.Series, keys: List[str], default: str = "") -> str:
    for k in keys:
        if k in row and pd.notna(row[k]) and str(row[k]).strip():
            return _clean(str(row[k]))
    return default

def parse_date(val) -> str:
    if val is None or (isinstance(val, str) and not val.strip()):
        return ""
    try:
        if isinstance(val, (pd.Timestamp, datetime, date)):
            dt = pd.to_datetime(val)
        else:
            dt = dtparse.parse(str(val), dayfirst=False, fuzzy=True)
        return dt.strftime("%m/%d/%Y")
    except Exception:
        return str(val)

def age_from_dob(dob: str) -> str:
    if not dob:
        return ""
    try:
        d = dtparse.parse(dob).date()
        today = date.today()
        years = today.year - d.year - ((today.month, today.day) < (d.month, d.day))
        return str(years)
    except Exception:
        return ""

def set_table_full_width(section, table, cols: int):
    avail = section.page_width - section.left_margin - section.right_margin
    table.allow_autofit = False
    colw = int(avail / cols)
    for i in range(cols):
        table.columns[i].width = colw

def _set_para_font(p, size_pt=FONT_SIZE_PT, bold=False, align=None):
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.size = Pt(size_pt)
        run.bold = bold

def _write_cell_lines(cell, lines: list[tuple[str, bool]]):
    cell.text = ""
    for text, bold in lines:
        p = cell.add_paragraph(text)
        _set_para_font(p, FONT_SIZE_PT, bold, align=None)

def cell_bottom_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    tcBorders.append(bottom)

# -------- document build --------

def build_doc(row: pd.Series, out_dir: str) -> str:
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)

    # Values
    caseno = coalesce(row, ["caseno"]) or "Unknown"
    visit_date = parse_date(coalesce(row, ["visitdate"]))
    visit_time = coalesce(row, ["visittime"]).strip()

    wardfirst = coalesce(row, ["wardfirst"]) or ""
    wardlast = coalesce(row, ["wardlast"]) or ""
    wname = _clean(f"{wardfirst} {wardlast}")
    wdob = parse_date(coalesce(row, ["wdob"]))
    wage = age_from_dob(wdob)
    wtele = coalesce(row, ["wtele"]) or ""
    waddr = coalesce(row, ["waddress"]) or ""
    liveswith = coalesce(row, ["liveswith"]) or ""

    g1 = coalesce(row, ["guardian1"]) or ""
    gdob = parse_date(coalesce(row, ["gdob"]))
    gtele = coalesce(row, ["gtele"]) or ""
    gemail = coalesce(row, ["gemail"]) or ""
    grel = coalesce(row, ["Relationship"]) or ""
    gaddr = coalesce(row, ["gaddress"]) or ""

    g2 = coalesce(row, ["Guardian2"]) or ""
    g2dob = parse_date(coalesce(row, ["g2dob"]))
    g2tele = coalesce(row, ["g2tele"]) or ""
    g2email = coalesce(row, ["g2eamil"]) or ""
    g2rel = coalesce(row, ["g2Relationship"]) or ""
    g2addr = coalesce(row, ["g2 address"]) or ""

    # Header line 1: small, bold
    title = doc.add_paragraph("Court Visitor Summary")
    _set_para_font(title, HEADER_FONT_SIZE_PT, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Header line 2: Visit (left) | Cause Number (right)
    hdr = doc.add_table(rows=1, cols=2)
    set_table_full_width(section, hdr, cols=2)
    hdr.style = None
    left = hdr.rows[0].cells[0]
    right = hdr.rows[0].cells[1]
    left.text = f"Visit: {visit_date} {(' ' + visit_time) if visit_time else ''}"
    right.text = f"Cause Number: {caseno}"
    for cell in (left, right):
        _set_para_font(cell.paragraphs[0], SUBHEAD_FONT_SIZE_PT, bold=False, align=None)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Top half info: three columns (Ward | Guardian 1 | Guardian 2)
    info = doc.add_table(rows=1, cols=3)
    set_table_full_width(section, info, cols=3)

    ward_lines = [
        (wname, True),
        (wtele, False),
        (waddr, False),
        (f"DOB: {wdob}" + (f"  |  Age: {wage}" if wage else ""), False),
        (f"Lives With: {liveswith}", False),
    ]
    _write_cell_lines(info.rows[0].cells[0], ward_lines)

    g1_lines = [
        (g1, True),
        (gtele, False),
        (gaddr, False),
        (gemail, False),
        (f"DOB: {gdob}" + (f"  |  {grel}" if grel else ""), False),
    ]
    _write_cell_lines(info.rows[0].cells[1], g1_lines)

    g2_lines = [
        (g2, True),
        (g2tele, False),
        (g2addr, False),
        (g2email, False),
        (f"DOB: {g2dob}" + (f"  |  {g2rel}" if g2rel else ""), False),
    ]
    _write_cell_lines(info.rows[0].cells[2], g2_lines)

    # Spacer
    doc.add_paragraph("")

    # Notes header
    nh = doc.add_paragraph("NOTES")
    _set_para_font(nh, FONT_SIZE_PT, bold=True, align=None)

    # Full-width ruled lines
    notes_table = doc.add_table(rows=NOTES_ROWS, cols=1)
    notes_table.allow_autofit = False
    avail = section.page_width - section.left_margin - section.right_margin
    notes_table.columns[0].width = avail
    for r in notes_table.rows:
        cell = r.cells[0]
        cell.text = "\u00A0"
        _set_para_font(cell.paragraphs[0], size_pt=10, bold=False, align=None)
        cell_bottom_border(cell)

    # Save
    os.makedirs(out_dir, exist_ok=True)
    safe_date = visit_date.replace('/', '-') if visit_date else ""
    filename = f"{caseno}_{wardlast}_{wardfirst}_{safe_date}.docx".replace('\\', '-').replace('/', '-')
    out_path = os.path.abspath(os.path.join(out_dir, filename))
    doc.save(out_path)
    return out_path

# -------- IO & CLI --------

def load_sheet_to_df(path: str, sheet: str | None) -> pd.DataFrame:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb.active
        rows = ws.values
        headers = next(rows)
        headers = [str(h).strip() if h is not None else "" for h in headers]
        df = pd.DataFrame(list(rows), columns=headers)
        return df
    finally:
        wb.close()

def pick_rows_gui(df: pd.DataFrame, last_n: int = 15) -> pd.DataFrame:
    import tkinter as tk
    # Use sheet order: last N physical rows, so newest entries always appear
    df2 = df.tail(last_n).copy()

    labels = []
    for _, r in df2.iterrows():
        labels.append(f"{r.get('caseno','')} — {r.get('wardlast','')}, {r.get('wardfirst','')} — {parse_date(r.get('visitdate',''))} {coalesce(r, ['visittime'])}")

    root = tk.Tk()
    root.title("Select Wards (Ctrl/Shift for multi)")
    root.geometry("900x420")

    lb = tk.Listbox(root, selectmode=tk.MULTIPLE, width=140, height=20)
    for s in labels:
        lb.insert(tk.END, s)
    lb.pack(fill=tk.BOTH, expand=True)

    selected = []
    def on_ok():
        nonlocal selected
        selected = list(lb.curselection())
        root.destroy()
    tk.Button(root, text="Generate", command=on_ok).pack(pady=8)

    root.mainloop()

    if not selected:
        return df2.iloc[0:0]
    return df2.iloc[selected]

def try_print_via_word(path: str) -> bool:
    if win32com is None:
        print("Printing requires pywin32. Install with:\n  py -3 -m pip install pywin32")
        return False
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(path)
        doc.PrintOut()
        doc.Close(False)
        word.Quit()
        return True
    except Exception as e:
        print(f"Could not print {path}: {e}")
        return False

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--workbook", default=WORKBOOK_PATH)
    ap.add_argument("--sheet", default=SHEET_NAME)
    ap.add_argument("--output", default=OUTPUT_DIR)
    ap.add_argument("--last", type=int, default=15, help="Show picker for last N (sheet order)")
    ap.add_argument("--no-gui", action="store_true", help="Disable picker; process all (or tail N) directly")
    ap.add_argument("--print", action="store_true", help="Send each generated docx to the default printer via Word")
    ap.add_argument("--open", action="store_true", help="Open the output folder in Explorer when done")
    args = ap.parse_args()

    if not os.path.exists(args.workbook):
        print(f"ERROR: Workbook not found: {args.workbook}")
        sys.exit(1)

    df = load_sheet_to_df(args.workbook, args.sheet)
    if df.empty:
        print("No data rows found.")
        sys.exit(1)
    df = normalize_columns(df)

    if args.no_gui:
        picked = df.tail(args.last) if args.last else df
    else:
        picked = pick_rows_gui(df, last_n=args.last)

    if picked is None or picked.empty:
        print("No rows selected.")
        return

    os.makedirs(args.output, exist_ok=True)
    for _, row in picked.iterrows():
        out_path = build_doc(row, args.output)
        print(f"Created: {out_path}")
        if args.print:
            try_print_via_word(out_path)

    print(f"All summaries saved to: {os.path.abspath(args.output)}")
    if args.open:
        try:
            folder_path = os.path.abspath(args.output)
            print(f"Opening folder: {folder_path}")
            os.startfile(folder_path)  # type: ignore[attr-defined]
        except Exception as e:
            print(f"WARNING: Could not open folder: {e}")

if __name__ == "__main__":
    import sys
    import traceback
    try:
        main()
        print("\n[OK] Visit Summary SUCCESS")
        sys.exit(0)
    except Exception as e:
        print(f"\n[FAIL] Visit Summary FAILED: {e}")
        traceback.print_exc()
        sys.exit(1)
