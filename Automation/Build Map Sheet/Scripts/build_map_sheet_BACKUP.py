# build_map_sheet.py
# Creates a printable map sheet from ward_guardian_info.xlsx
# Page 1: basemap with numbered markers + legend (#, wardlast, address)
# Page 2 (optional): suggested route if visitdate & visittime exist

import os
import math
import time
from datetime import datetime
from dateutil import parser as dtparser

import pandas as pd
from geopy.geocoders import Nominatim
from geopy.extra.rate_limiter import RateLimiter

import matplotlib.pyplot as plt
try:
    import contextily as cx
    HAS_CONTEXTILY = True
except ImportError:
    HAS_CONTEXTILY = False
    print("WARNING: contextily not installed. Maps will be created without basemap tiles.")
from pyproj import Transformer
from io import BytesIO
import requests

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------
# CONFIG — EDIT THESE PATHS
# ----------------------------
# Use the Excel file from the app directory (3 levels up from this script)
EXCEL_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "..", "ward_guardian_info.xlsx")
# Output to the app directory
OUTPUT_DOCX = os.path.join(os.path.dirname(__file__), "..", "..", "..", "Ward_Map_Sheet.docx")

SHEET_NAME = 0  # or 'Sheet1'

# Column names in your workbook (adjust if different)
COL_WARDLAST     = "wardlast"
COL_ADDRESS      = "waddress"
COL_DATE_SUBMIT  = "datesubmitted"
COL_VISIT_DATE   = "visitdate"
COL_VISIT_TIME   = "visittime"


# Optional: Treat blank-like values
BLANK_LIKE = {"", " ", "NA", "N/A", "nan", "none", None}

# Geocoding
GEOCODER_USER_AGENT = "guardian-map-sheet"
GEOCODE_SLEEP_SEC = 1.0  # polite delay for Nominatim
# Cache for geocoding results.  We store it alongside the script folder so it
# persists across runs.  The directory is created on demand when saving.
# Cache in the app directory
GEOCODE_CACHE_CSV = os.path.join(os.path.dirname(__file__), "..", "..", "..", "geocode_cache.csv")

# Routing options
# Pull the user-provided Google key from the environment and strip whitespace.
# A non-empty key means we can use Google for both geocoding and directions.
GOOGLE_MAPS_API_KEY = os.environ.get("GOOGLE_MAPS_API_KEY", "").strip()
# Flags for enabling Google services.  We expose separate booleans so the rest
# of the script can check them without having to know how the key is stored.
USE_GOOGLE_GEOCODER: bool = bool(GOOGLE_MAPS_API_KEY)
USE_GOOGLE_DIRECTIONS: bool = bool(GOOGLE_MAPS_API_KEY)

# OSRM fallback for ETA if Google not used (public demo server; best-effort)
OSRM_ROUTE_URL = "https://router.project-osrm.org/route/v1/driving/{lonlat_list}?overview=false"

# --- A) Auto-open the finished DOCX ---
OPEN_ON_COMPLETE = True

# ----------------------------
# Helpers
# ----------------------------
def is_blank(x):
    # Treat None/NaN/NaT and empty-ish strings as blank
    try:
        import pandas as pd
    except Exception:
        pd = None

    if x is None:
        return True
    if pd is not None and pd.isna(x):   # catches NaN and NaT
        return True

    s = str(x).replace("\u00A0", " ").replace("\u200B", "").strip().lower()
    return s in {"", "na", "n/a", "nan", "none", "nat"}


# -----------------------------------------------------------------------------
# Column resolution and address normalisation helpers
# -----------------------------------------------------------------------------

# We support multiple column names for the same logical field.  If a user
# renames a column (for example from "wardlast" to "lastname"), the script
# attempts to pick the first match from these alias lists.
ADDRESS_ALIASES    = ["waddress", "wardaddress", "address"]
WARDLAST_ALIASES   = ["wardlast", "ward_last", "last", "lastname"]
DATESUBMIT_ALIASES = ["datesubmitted", "date_submitted", "submitted"]
VISITDATE_ALIASES  = ["visitdate", "visit_date", "datevisit", "date_visit"]
VISITTIME_ALIASES  = ["visittime", "visit_time", "timevisit", "time_visit"]

def pick_col(df: pd.DataFrame, aliases, required: bool = True, label: str = "column"):
    """Return the first column in df that matches one of the aliases.

    If none of the aliases exist and required is True, raise an error; otherwise
    return None.  This centralises the logic for resolving flexible column names.
    """
    for a in aliases:
        if a in df.columns:
            return a
    if required:
        raise ValueError(f"Missing required {label}. Tried aliases: {aliases}")
    return None


def _normalize_addr(a: str) -> str:
    """Normalise an address string before geocoding.

    Collapses multiple spaces, strips whitespace, and converts slash-separated
    components like "Austin/TX/78734" into a format accepted by most geocoders
    ("Austin, TX 78734").  Returns the original string if no normalisation
    applies.
    """
    if not a:
        return a
    s = str(a).strip()
    # Convert "City/ST/ZIP" into "City, ST ZIP" if no comma is present
    if "/" in s and "," not in s:
        parts = [p.strip() for p in s.split("/") if p.strip()]
        if len(parts) >= 2:
            # If more than two parts, treat all leading parts except the last two as the city
            if len(parts) > 2:
                city = " ".join(parts[:-2])
                state = parts[-2]
                zipc  = parts[-1]
            else:
                city = parts[0]
                state = parts[1]
                zipc  = parts[2] if len(parts) > 2 else ""
            s = f"{city}, {state} {zipc}".strip()
    # Collapse duplicate spaces
    s = " ".join(s.split())
    return s


def load_geocode_cache(path):
    if os.path.exists(path):
        try:
            df = pd.read_csv(path)
            return df
        except Exception:
            return pd.DataFrame(columns=["address", "lat", "lon"])
    return pd.DataFrame(columns=["address", "lat", "lon"])

def save_geocode_cache(df, path):
    cache_dir = os.path.dirname(path) or "."
    os.makedirs(cache_dir, exist_ok=True)
    df.drop_duplicates(subset=["address"], keep="last").to_csv(path, index=False)

 
def geocode_addresses(df):
    """Geocode the addresses in ``df`` using either Google or Nominatim.

    This function respects the ``USE_GOOGLE_GEOCODER`` flag.  When True and a
    Google API key is present, it will attempt to use the Google Geocoding
    API.  Otherwise it falls back to Nominatim with a rate limiter.  Results
    are cached on disk so repeated runs are fast and do not repeatedly hit
    external services.  The cache key is the normalised address string.

    For each address we log whether Google or Nominatim was used and whether
    geocoding succeeded.  At the end we print a summary of how many lookups
    succeeded and how many failed, and list the failed addresses.
    """
    # Load existing cache
    cache = load_geocode_cache(GEOCODE_CACHE_CSV)
    cached = {a: (lat, lon) for a, lat, lon in zip(cache["address"], cache["lat"], cache["lon"])}

    # Prepare Google client if enabled
    gmaps = None
    if USE_GOOGLE_GEOCODER:
        try:
            import googlemaps  # type: ignore
            gmaps = googlemaps.Client(key=GOOGLE_MAPS_API_KEY)
        except Exception as e:
            # If we cannot initialise Google client, disable Google for this run
            print(f"[geocode] Could not initialise Google geocoder: {e}")
            gmaps = None

    # Prepare Nominatim fallback with retries and timeout
    nominatim = None
    rate_limited = None
    if not USE_GOOGLE_GEOCODER or gmaps is None:
        # Only set up Nominatim if we aren't using Google or if Google init failed
        try:
            nominatim = Nominatim(user_agent=GEOCODER_USER_AGENT, timeout=10)
            rate_limited = RateLimiter(
                nominatim.geocode,
                min_delay_seconds=GEOCODE_SLEEP_SEC,
                max_retries=3,
                error_wait_seconds=2.0,
                swallow_exceptions=True,
            )
        except Exception as e:
            print(f"[geocode] Could not initialise Nominatim: {e}")
            nominatim = None
            rate_limited = None

    lats: list[float | None] = []
    lons: list[float | None] = []
    failed: list[str] = []

    # Iterate through each address in the DataFrame
    for raw in df[COL_ADDRESS]:
        # Normalise and trim the address for consistent caching
        addr_str = str(raw).strip()
        addr_key = _normalize_addr(addr_str)

        if not addr_key:
            # Skip empty strings entirely
            lats.append(None)
            lons.append(None)
            failed.append(addr_str)
            continue

        # Use cached result if available
        if addr_key in cached:
            lat, lon = cached[addr_key]
            lats.append(lat)
            lons.append(lon)
            continue

        lat = None
        lon = None

        # Try Google geocoding first if configured
        if gmaps is not None:
            try:
                geos = gmaps.geocode(addr_key)
                if geos:
                    loc = geos[0]["geometry"]["location"]
                    lat = loc["lat"]
                    lon = loc["lng"]
                    print(f"[geocode] Google OK: {addr_key} -> ({lat:.6f},{lon:.6f})")
                else:
                    print(f"[geocode] Google returned 0 results for: {addr_key}")
            except Exception as e:
                # Log the error and fall back; do not include the key
                print(f"[geocode] Google error for '{addr_key}': {e}")

        # If Google failed or was not used, try Nominatim
        if (lat is None or lon is None) and rate_limited is not None:
            try:
                loc = rate_limited(addr_key)
                if loc:
                    lat = loc.latitude
                    lon = loc.longitude
                    print(f"[geocode] Nominatim OK: {addr_key} -> ({lat:.6f},{lon:.6f})")
                else:
                    print(f"[geocode] Nominatim returned 0 results for: {addr_key}")
            except Exception as e:
                print(f"[geocode] Nominatim error for '{addr_key}': {e}")

        # Record results and update cache if successful
        if lat is None or lon is None:
            lats.append(None)
            lons.append(None)
            failed.append(addr_str)
        else:
            lats.append(lat)
            lons.append(lon)
            cached[addr_key] = (lat, lon)

    # Persist the updated cache
    new_rows: list[dict[str, float | str]] = []
    for addr_raw, lat, lon in zip(df[COL_ADDRESS], lats, lons):
        if lat is not None and lon is not None:
            new_rows.append({"address": _normalize_addr(str(addr_raw).strip()), "lat": lat, "lon": lon})
    if new_rows:
        new_cache = pd.concat([cache, pd.DataFrame(new_rows)], ignore_index=True)
        save_geocode_cache(new_cache, GEOCODE_CACHE_CSV)

    # Report summary
    success_count = sum(1 for lat in lats if lat is not None)
    fail_count = len(failed)
    print(f"[geocode] success={success_count} failed={fail_count}")
    if failed:
        for a in failed:
            print(f"   - {a}")

    # Attach lat/lon columns onto a copy of the input DataFrame
    df_out = df.copy()
    df_out["lat"] = lats
    df_out["lon"] = lons
    return df_out, failed

def to_web_mercator(lats, lons):
    transformer = Transformer.from_crs("epsg:4326", "epsg:3857", always_xy=True)
    xs, ys = transformer.transform(lons.tolist(), lats.tolist())
    return xs, ys

def build_map_image(filtered_df, dpi=200):
    # Prepare data
    plot_df = filtered_df.dropna(subset=["lat", "lon"]).reset_index(drop=True)
    if plot_df.empty:
        return None  # nothing to plot

    xs, ys = to_web_mercator(plot_df["lat"], plot_df["lon"])
    plot_df["x"] = xs; plot_df["y"] = ys

    # Determine map extent with padding
    pad = 500  # meters
    xmin, xmax = plot_df["x"].min()-pad, plot_df["x"].max()+pad
    ymin, ymax = plot_df["y"].min()-pad, plot_df["y"].max()+pad

    # Plot
    fig = plt.figure(figsize=(8.0, 6.0), dpi=dpi)
    ax = plt.gca()
    ax.set_xlim([xmin, xmax]); ax.set_ylim([ymin, ymax])

    # Basemap wrapped safely
    if HAS_CONTEXTILY:
        try:
            cx.add_basemap(ax, crs="EPSG:3857")  # OSM tiles
        except Exception as e:
            print(f"Warning: basemap unavailable ({e}); proceeding without tiles.")
    else:
        print("Info: Proceeding without basemap tiles (contextily not installed).")

    # Numbered markers
    for idx, row in plot_df.iterrows():
        num = idx + 1
        ax.plot(row["x"], row["y"], marker="o", markersize=6, linewidth=0)
        ax.text(row["x"] + 5, row["y"] + 5, str(num), fontsize=9, weight="bold")

    ax.set_axis_off()
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)
    buf.seek(0)
    return buf


def parse_visit_dt(row):
    """Combine visitdate + visittime to a single datetime (best effort)."""
    d, t = row.get(COL_VISIT_DATE), row.get(COL_VISIT_TIME)
    if is_blank(d) or is_blank(t):
        return None
    try:
        # Allow Excel date/time formats or strings
        if not isinstance(d, (str, datetime)):
            # pandas timestamp -> python dt
            d = pd.to_datetime(d).to_pydatetime()
        else:
            d = dtparser.parse(str(d), fuzzy=True)

        if not isinstance(t, (str, datetime)):
            t = pd.to_datetime(t).to_pydatetime().time()
        else:
            t = dtparser.parse(str(t), fuzzy=True).time()

        return datetime.combine(d.date(), t)
    except Exception:
        return None


def estimate_durations(points):
    """Estimate travel durations between consecutive points.
       Tries Google Directions if enabled; else uses OSRM; else haversine fallback.
       Returns list of (leg_minutes)."""
    legs_min = []

    def haversine_minutes(lat1, lon1, lat2, lon2, avg_mph=28.0):
        # 28 mph ~ city driving average incl. stops
        R = 3958.8  # miles
        phi1 = math.radians(lat1); phi2 = math.radians(lat2)
        dphi = math.radians(lat2 - lat1)
        dlambda = math.radians(lon2 - lon1)
        a = math.sin(dphi/2)**2 + math.cos(phi1)*math.cos(phi2)*math.sin(dlambda/2)**2
        dist = 2*R*math.asin(math.sqrt(a))
        hours = dist / max(avg_mph, 1e-6)
        return round(hours*60)

    if USE_GOOGLE_DIRECTIONS:
        try:
            import googlemaps  # type: ignore
            gmaps = googlemaps.Client(key=GOOGLE_MAPS_API_KEY)
            for i in range(len(points)-1):
                o = points[i]
                d = points[i+1]
                try:
                    res = gmaps.directions((o["lat"], o["lon"]), (d["lat"], d["lon"]), mode="driving")
                    if res and "legs" in res[0] and res[0]["legs"]:
                        secs = res[0]["legs"][0]["duration"]["value"]
                        legs_min.append(round(secs/60))
                    else:
                        legs_min.append(haversine_minutes(o["lat"], o["lon"], d["lat"], d["lon"]))
                except Exception:
                    # Fallback to straight-line estimate for this leg
                    legs_min.append(haversine_minutes(o["lat"], o["lon"], d["lat"], d["lon"]))
            return legs_min
        except Exception as e:
            # If Google Directions is not available, we'll fall back to OSRM/haversine below
            pass

    # Try OSRM public API
    try:
        lonlat_list = ";".join([f'{p["lon"]},{p["lat"]}' for p in points])
        url = OSRM_ROUTE_URL.format(lonlat_list=lonlat_list)
        r = requests.get(url, timeout=10)
        if r.ok:
            data = r.json()
            if data.get("routes"):
                # OSRM returns total duration only; approximate per leg by splitting equally by haversine weights
                total_sec = data["routes"][0]["duration"]
                # weight by haversine distances
                dists = []
                for i in range(len(points)-1):
                    dists.append(
                        max(0.01, math.dist(
                            [points[i]["lat"], points[i]["lon"]],
                            [points[i+1]["lat"], points[i+1]["lon"]]
                        ))
                    )
                s = sum(dists)
                if s > 0:
                    legs_min = [round((total_sec/60)*(w/s)) for w in dists]
                else:
                    legs_min = [5]*(len(points)-1)
                return legs_min
    except Exception:
        pass

    # Fallback
    for i in range(len(points)-1):
        o, d = points[i], points[i+1]
        legs_min.append(haversine_minutes(o["lat"], o["lon"], d["lat"], d["lon"]))
    return legs_min

def make_doc(map_image_buf, legend_rows, route_rows):
    doc = Document()

    # Title
    doc.add_heading("Ward Visit Map", level=1)

    # Map image
    if map_image_buf is not None:
        pic = doc.add_picture(map_image_buf, width=Inches(6.5))
        last_par = doc.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("No mappable addresses found.")
        p.runs[0].bold = True

    # Legend
    doc.add_heading("Legend", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Ward Last"
    hdr[2].text = "Address"
    for row in legend_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["num"])
        cells[1].text = str(row["wardlast"])
        cells[2].text = str(row["address"])

    # Page 2: Route suggestion (if any)
    if route_rows:
        doc.add_page_break()
        doc.add_heading("Suggested Route (by visit date/time)", level=1)
        table2 = doc.add_table(rows=1, cols=6)
        h2 = table2.rows[0].cells
        h2[0].text = "#"
        h2[1].text = "When"
        h2[2].text = "Ward Last"
        h2[3].text = "Address"
        h2[4].text = "Drive (min)"
        h2[5].text = "Arrive By"

        for rr in route_rows:
            r = table2.add_row().cells
            r[0].text = str(rr["num"])
            r[1].text = rr["when"]
            r[2].text = rr["wardlast"]
            r[3].text = rr["address"]
            r[4].text = rr["drive_min"]
            r[5].text = rr["arrive_by"]

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)

# ----------------------------
# Main
# ----------------------------
def main():
    # Load the spreadsheet
    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)



    # Report whether a Google key was detected so the user knows which geocoder will be used
    print(f"[init] Google key detected: {USE_GOOGLE_GEOCODER}; Directions: {USE_GOOGLE_DIRECTIONS}")

    # Resolve column names from aliases.  This allows minor header changes without
    # breaking the script.  Required columns must exist; optional columns may be
    # missing and will be treated as blank.
    col_address    = pick_col(df, ADDRESS_ALIASES, label="address column")
    col_wardlast   = pick_col(df, WARDLAST_ALIASES, label="ward last name column")
    col_datesubmit = pick_col(df, DATESUBMIT_ALIASES, label="datesubmitted column")
    col_visitdate  = pick_col(df, VISITDATE_ALIASES, required=False, label="visitdate column")
    col_visittime  = pick_col(df, VISITTIME_ALIASES, required=False, label="visittime column")

    # Update global constants so helper functions see the resolved names
    global COL_ADDRESS, COL_WARDLAST, COL_DATE_SUBMIT, COL_VISIT_DATE, COL_VISIT_TIME
    COL_ADDRESS = col_address
    COL_WARDLAST = col_wardlast
    COL_DATE_SUBMIT = col_datesubmit
    # For optional columns, only update if they were found; otherwise leave defaults
    if col_visitdate:
        COL_VISIT_DATE = col_visitdate
    if col_visittime:
        COL_VISIT_TIME = col_visittime

    # Filter rows where the datesubmitted column is blank
    filt = df[df[COL_DATE_SUBMIT].apply(is_blank)].copy()

    # Require at least the address and ward name columns to exist
    required_cols = [COL_WARDLAST, COL_ADDRESS, COL_DATE_SUBMIT]
    for c in required_cols:
        if c not in df.columns:
            raise ValueError(f"Missing required column: {c}")

    # Clean addresses and drop rows with blank addresses
    filt[COL_ADDRESS] = filt[COL_ADDRESS].astype(str).str.strip()
    filt = filt[~filt[COL_ADDRESS].apply(is_blank)].reset_index(drop=True)

    if filt.empty:
        print("No rows where datesubmitted is blank and address present.")
        # Always write a document even if empty so downstream processes don't fail
        print(f"[output] writing: {OUTPUT_DOCX}")
        make_doc(None, [], [])
        print(f"Map sheet created: {OUTPUT_DOCX}")
        # --- A) Auto-open the finished DOCX ---
        if OPEN_ON_COMPLETE:
            try:
                os.startfile(OUTPUT_DOCX)
            except Exception as e:
                print(f"[warn] Could not open file automatically: {e}")
        return

    # Geocode addresses
    filt, failed = geocode_addresses(filt)
    if failed:
        # Already logged individually; just summarise count here
        print(f"[geocode] {len(failed)} address(es) could not be geocoded and will be skipped on the map.")

    # Only keep rows where geocoding succeeded
    plot_df = filt.dropna(subset=["lat", "lon"]).reset_index(drop=True)

    # Assign sequential numbers for the legend and map markers
    plot_df["num"] = range(1, len(plot_df) + 1)

    # Build the map image
    map_img = build_map_image(plot_df)

    # Prepare legend rows for the Word table
    legend_rows: list[dict] = []
    for _, r in plot_df.iterrows():
        legend_rows.append({
            "num": int(r["num"]),
            "wardlast": r[COL_WARDLAST],
            "address": r[COL_ADDRESS],
        })

    # Prepare route rows: include only rows with both visit date and time
    route_df = plot_df.copy()
    route_df["visit_dt"] = route_df.apply(parse_visit_dt, axis=1)
    route_df = route_df.dropna(subset=["visit_dt"]).sort_values("visit_dt").reset_index(drop=True)

    route_rows: list[dict] = []
    if not route_df.empty:
        points = [{"lat": r["lat"], "lon": r["lon"]} for _, r in route_df.iterrows()]
        legs_min = estimate_durations(points) if len(points) >= 2 else []
        for i, r in route_df.iterrows():
            num = i + 1
            when_str = r["visit_dt"].strftime("%Y-%m-%d %I:%M %p")
            drive_min = ""
            if i > 0 and (i - 1) < len(legs_min):
                drive_min = str(legs_min[i - 1])
            route_rows.append({
                "num": num,
                "when": when_str,
                "wardlast": str(r[COL_WARDLAST]),
                "address": str(r[COL_ADDRESS]),
                "drive_min": drive_min,
                "arrive_by": when_str,
            })

    # Write the Word document.  Log the write operation before invoking the helper.
    print(f"[output] writing: {OUTPUT_DOCX}")
    make_doc(map_img, legend_rows, route_rows)
    print(f"Map sheet created: {OUTPUT_DOCX}")

    # --- A) Auto-open the finished DOCX ---
    if OPEN_ON_COMPLETE:
        try:
            os.startfile(OUTPUT_DOCX)  # Windows
        except Exception as e:
            print(f"[warn] Could not open file automatically: {e}")

if __name__ == "__main__":
    main()
