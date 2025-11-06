"""
Court Visitor Payment form auto-filler (v3)
- Reads Excel at EXCEL_PATH (read-only; never writes back)
- Prompts for month (supports many formats; see parse_month_input)
- Filters to that month (optionally month-to-date)
- Auto-selects the template table by header text; does NOT add rows
- Fills exactly 7 lines per form (1 header + 7 data rows expected)
- Saves outputs as "9_1Sept2025 Court Visitor Payment.docx", etc., in OUTPUT_DIR

Dependencies (once):
  python -m pip install pandas python-docx python-dateutil openpyxl
"""
from __future__ import annotations
import os
import re
from dataclasses import dataclass
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
from typing import List, Dict, Any, Optional, Sequence

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ===========================
# CONFIG — EDIT THESE IF NEEDED
# ===========================
EXCEL_PATH = r"C:\Users\may\OneDrive\Guardian Docs\ward_guardian_info.xlsx"  # READ-ONLY
SHEET_NAME = 0  # first sheet
DATE_COLUMN = "visitdate"  # used to select the month
OUTPUT_DIR = r"G:\My Drive\Guardianship files\Payments to review and submit"
TEMPLATE_PATH = r"C:\GuardianAutomation\2025 Court Visitor Payment Invoice - fields.docx"  # point to your local .docx

# WORD HEADER TEXT we expect (case-insensitive "contains" match; used to pick the correct table)
EXPECTED_HEADER_BITS = [
    "cause",                      # Cause No.
    "date appointed",
    "date of court visit",
    "date of court visitor report",
]

# The 4 table columns mapping (left→right) to Excel headers. We’ll resolve aliases.
COLUMN_MAP_BASE: Dict[int, Sequence[str]] = {
    0: ("causeno",),                              # Cause No.
    1: ("Dateappointed", "dateappointed"),        # Date Appointed
    2: ("visitdate",),                            # Date of Court Visit
    3: ("datesubmitted", "datesubmited"),         # Date of Court Visitor Report (accepts both spellings)
}

# Date formatting on the form
DATE_FORMAT = "%m/%d/%Y"  # e.g., 09/05/2025

# Output file name pattern
FILENAME_PATTERN = "{BILL_MONTH_NUM}_{FORM_NO}{BILL_MONTH_NAME_SHORT}{BILL_YEAR} Court Visitor Payment.docx"

# ===========================
# END CONFIG
# ===========================

def month_short_name(dt: date) -> str:
    """Return short month name with 'Sept' for September (instead of 'Sep')."""
    name = dt.strftime("%b")
    return "Sept" if name == "Sep" else name

# --- Table detection: pick the 7-line table with correct header ---
EXPECTED_HEADER_BITS = [
    "cause",
    "date appointed",
    "date of court visit",
    "date of court visitor report",
]

def find_target_table(doc: Document) -> int:
    """
    Choose the table that:
      - has >= 8 rows (1 header + 7 data rows),
      - has >= 4 header cells,
      - header contains the expected labels (case-insensitive).
    Fallback: the table with the most rows.
    """
    best_idx = None
    best_score = -1
    fallback_idx = None
    fallback_rows = -1
    for idx, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.rows[0].cells) if rows else 0
        if rows > fallback_rows:
            fallback_rows = rows
            fallback_idx = idx
        if rows < 8 or cols < 4:
            continue
        header_text = " | ".join(c.text.strip().lower() for c in t.rows[0].cells)
        score = sum(bit in header_text for bit in EXPECTED_HEADER_BITS)
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx if best_idx is not None else (fallback_idx if fallback_idx is not None else 0)

# --- Cause-number helper: keep the template's "C-1-PB-" and only insert the suffix ---
def cause_suffix_from_value(raw) -> str:
    import pandas as pd, re
    if pd.isna(raw):
        return ""
    s = str(raw).strip()
    # strip any "C-1-PB" variant at the start
    return re.sub(r"^\s*(?i:c-?1-?pb-?)\s*", "", s)

# --- Date formatting helper ---
def format_date_like(val, out_fmt="%m/%d/%Y") -> str:
    import pandas as pd
    from datetime import date, datetime
    if pd.isna(val):
        return ""
    try:
        if isinstance(val, date):
            return val.strftime(out_fmt)
        if isinstance(val, datetime):
            return val.date().strftime(out_fmt)
        dt = pd.to_datetime(val, errors="coerce")
        if pd.notna(dt):
            return dt.date().strftime(out_fmt)
    except Exception:
        pass
    return str(val)

# --- Low-level SDT writer (content controls) ---
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

def _set_sdt_text(sdt_elm, new_text: str):
    """
    Write plain text into an SDT's <w:sdtContent>.
    If there is no <w:t>, create one. Overwrite existing text runs.
    """
    if new_text is None:
        new_text = ""
    sdt_content = sdt_elm.find(".//w:sdtContent", NS)
    if sdt_content is None:
        return
    texts = sdt_content.findall(".//w:t", NS)
    if texts:
        texts[-1].text = new_text
        return
    # Create <w:p><w:r><w:t> if missing
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = new_text
    r.append(t)
    p.append(r)
    sdt_content.append(p)

# Pick the first header name that actually exists in df.columns (case-insensitive)
def first_existing_header(df, *names):
    lowers = {c.lower(): c for c in df.columns}
    for n in names:
        if n in df.columns:
            return n
        if n.lower() in lowers:
            return lowers[n.lower()]
    # fallback (shouldn't happen if headers are present)
    return names[0]


def parse_month_input(raw: str) -> Optional[date]:
    """
    Accepts many formats:
      - 'last' (previous month)
      - '' (blank) -> current month
      - 'YYYY-MM', 'YYYY/MM'
      - 'MM-YYYY', 'MM/YYYY'
      - 'M-YYYY', 'M/YYYY'
      - 'YYYY-M', 'YYYY/M'
      - '9', '09'  (uses current year)
    Returns date set to the FIRST of the target month.
    """
    raw = (raw or "").strip().lower()
    today = date.today()

    if raw == "":
        return today.replace(day=1)
    if raw == "last":
        return (today.replace(day=1) - relativedelta(months=1))

    # Digit-only month like '9' or '09'
    if raw.isdigit():
        m = int(raw)
        if 1 <= m <= 12:
            return date(today.year, m, 1)

    # Try multiple explicit formats
    candidates = [
        ("%Y-%m", raw),
        ("%Y/%m", raw),
        ("%m-%Y", raw),
        ("%m/%Y", raw),
    ]

    # Also handle 'YYYY-M' or 'M-YYYY' where month may be 1 digit using regex
    # Patterns to normalize like '2025-9' -> '2025-09'
    m1 = re.fullmatch(r"(\d{4})[-/](\d{1,2})", raw)
    m2 = re.fullmatch(r"(\d{1,2})[-/](\d{4})", raw)
    if m1:
        y, m = int(m1.group(1)), int(m1.group(2))
        if 1 <= m <= 12:
            return date(y, m, 1)
    if m2:
        m, y = int(m2.group(1)), int(m2.group(2))
        if 1 <= m <= 12:
            return date(y, m, 1)

    # Try strict strptime attempts
    for fmt, text in candidates:
        try:
            dt = datetime.strptime(text, fmt).date()
            return dt.replace(day=1)
        except ValueError:
            pass

    # Could not parse
    return None

def prompt_for_month() -> (date, bool):
    """Ask user which month to compile. Returns (anchor_date_first_of_month, use_month_to_date)."""
    raw = input(
        "Enter month to compile (e.g., 9/2025, 09-2025, 2025-9, 2025/9, YYYY-MM). "
        "Blank = current month, or 'last' = previous month: "
    )
    anchor = parse_month_input(raw)
    if anchor is None:
        print("Could not parse; defaulting to current month.")
        anchor = date.today().replace(day=1)

    today = date.today()
    use_mtd = False
    if anchor.year == today.year and anchor.month == today.month:
        yn = input("Use month-to-date for the current month? [Y/n]: ").strip().lower()
        use_mtd = (yn in ("", "y", "yes"))
    else:
        # For non-current months, default to FULL month unless user says otherwise
        yn = input("Compile the FULL month (not capped to today)? [Y/n]: ").strip().lower()
        use_mtd = not (yn in ("", "y", "yes"))  # if YES full month -> not MTD
    return anchor, use_mtd

@dataclass
class BillingWindow:
    start: date
    end: date  # exclusive

    @staticmethod
    def for_month(anchor: date, to_date: Optional[date] = None) -> "BillingWindow":
        start = anchor.replace(day=1)
        end = (start + relativedelta(months=1))
        if to_date is not None:
            end = min(end, to_date + timedelta(days=1))  # half-open interval
        return BillingWindow(start, end)

def ensure_output_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def first_existing_col(df: pd.DataFrame, candidates: Sequence[str]) -> Optional[str]:
    # exact
    for c in candidates:
        if c in df.columns:
            return c
    # case-insensitive
    lower_map = {c.lower(): c for c in df.columns}
    for c in candidates:
        lc = c.lower()
        if lc in lower_map:
            return lower_map[lc]
    return None

def load_rows_for_window(bill_start: date, use_month_to_date: bool) -> pd.DataFrame:
    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, engine="openpyxl")

    if DATE_COLUMN not in df.columns:
        raise ValueError(f"Expected date column '{DATE_COLUMN}' not found in Excel. Found: {list(df.columns)}")

    df[DATE_COLUMN] = pd.to_datetime(df[DATE_COLUMN], errors="coerce").dt.date

    to_date_limit = date.today() if use_month_to_date else None
    window = BillingWindow.for_month(bill_start, to_date=to_date_limit)
    mask = (df[DATE_COLUMN] >= window.start) & (df[DATE_COLUMN] < window.end)
    month_df = df.loc[mask].copy()

    # Sort by visitdate then cause number if present
    sort_keys = [DATE_COLUMN]
    if "causeno" in month_df.columns:
        sort_keys.append("causeno")
    month_df.sort_values(sort_keys, inplace=True)

    # Resolve column aliases
    resolved_map: Dict[int, str] = {}
    for col_idx, aliases in COLUMN_MAP_BASE.items():
        found = first_existing_col(month_df, aliases)
        if not found:
            raise ValueError(
                f"None of the aliases {aliases} exist in the Excel header. "
                f"Headers present: {list(month_df.columns)}"
            )
        resolved_map[col_idx] = found
    month_df.attrs["resolved_map"] = resolved_map

    return month_df.reset_index(drop=True)

def find_target_table(doc: Document) -> int:
    """
    Pick the table that:
      - has at least 8 rows (1 header + 7 data rows),
      - has at least 4 columns,
      - whose header row contains our expected header bits (case-insensitive),
    Falling back to the table with the most rows if no header match is found.
    """
    best_idx = None
    best_score = -1
    fallback_idx = None
    fallback_rows = -1

    for idx, t in enumerate(doc.tables):
        row_count = len(t.rows)
        col_count = len(t.columns) if row_count else 0

        # Track a fallback (largest table)
        if row_count > fallback_rows:
            fallback_rows = row_count
            fallback_idx = idx

        if row_count >= 8 and col_count >= 4:
            try:
                header_text = " | ".join(cell.text.strip() for cell in t.rows[0].cells).lower()
            except Exception:
                header_text = ""
            score = sum(bit in header_text for bit in EXPECTED_HEADER_BITS)
            if score > best_score:
                best_score = score
                best_idx = idx

    if best_idx is not None:
        return best_idx
    # No clear header match; use the largest table as a fallback
    return fallback_idx if fallback_idx is not None else 0

def format_cell_value(excel_col_name: str, val: Any) -> str:
    if pd.isna(val):
        return ""
    # If the column appears to be a date field, try to format
    if any(k in excel_col_name.lower() for k in ["date", "appointed", "submitted", "submited"]):
        try:
            if isinstance(val, date):
                return val.strftime(DATE_FORMAT)
            if isinstance(val, datetime):
                return val.date().strftime(DATE_FORMAT)
            parsed = pd.to_datetime(val, errors="coerce")
            if pd.notna(parsed):
                return parsed.date().strftime(DATE_FORMAT)
        except Exception:
            pass
    return str(val)

def clear_table_data_area(t) -> None:
    """Clear rows 1..7 (keep header at row 0)."""
    # Do NOT change structure, only clear existing text.
    max_data_rows = min(len(t.rows) - 1, 7)
    for r in range(1, 1 + max_data_rows):
        for c in range(len(t.columns)):
            try:
                cell = t.cell(r, c)
                # clear all paragraphs in the cell
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                if not cell.paragraphs:
                    cell.add_paragraph("")
            except Exception:
                pass

def fill_table(doc: Document, records: List[Dict[str, Any]], resolved_map: Dict[int, str], table_index: int) -> None:
    """
    Fill the 7 data rows by writing into SDTs tagged:
      - causeno
      - dateappointed
      - visitdate
      - datesubmitted / datesubmited
    Does NOT alter the table structure; only populates the SDTs.
    """
    tables = doc.tables
    if not tables:
        raise RuntimeError("Template has no tables.")
    try:
        t = tables[table_index]
    except IndexError:
        raise RuntimeError(f"Template does not have table index {table_index}; found {len(tables)} tables.")

    required_rows = 8  # 1 header + 7 data rows
    if len(t.rows) < required_rows:
        raise RuntimeError(f"Selected table has {len(t.rows)} rows; need at least {required_rows} (header + 7).")

    # Build logical names from your resolved_map (index -> excel_col)
    # Expecting: 0:cause, 1:date appointed, 2:visitdate, 3:datesubmitted/datesubmited
    logical_from_excel = {
        "cause": resolved_map.get(0),
        "dateappointed": resolved_map.get(1),
        "visitdate": resolved_map.get(2),
        # accept either spelling for "date of report"
        "datesubmitted": resolved_map.get(3) if resolved_map.get(3) in ("datesubmitted", "datesubmited") else "datesubmitted",
    }

    # For the 7 data rows (row indices 1..7)
    for row_idx, record in enumerate(records[:7], start=1):
        tr = t.rows[row_idx]._tr  # OOXML row
        sdts = tr.xpath(".//w:sdt", namespaces=NS)

        # Map SDT tag -> SDT element for this row
        row_sdts = {}
        for sdt in sdts:
            tag = sdt.find(".//w:tag", NS)
            tag_val = tag.get(qn("w:val")) if tag is not None else None
            if tag_val:
                row_sdts[tag_val.lower()] = sdt

        # Cause suffix into the 'causeno' SDT (keeps the 'C-1-PB-' that's already in the template cell)
        cause_col = logical_from_excel.get("cause")
        if cause_col and "causeno" in row_sdts:
            suffix = cause_suffix_from_value(record.get(cause_col, ""))
            _set_sdt_text(row_sdts["causeno"], suffix)

        # Date appointed
        da_col = logical_from_excel.get("dateappointed")
        if da_col and "dateappointed" in row_sdts:
            _set_sdt_text(row_sdts["dateappointed"], format_date_like(record.get(da_col, "")))

        # Date of visit
        dv_col = logical_from_excel.get("visitdate")
        if dv_col and "visitdate" in row_sdts:
            _set_sdt_text(row_sdts["visitdate"], format_date_like(record.get(dv_col, "")))

        # Date of report (tag may be 'datesubmitted' OR 'datesubmited' in the template)
        dr_col = logical_from_excel.get("datesubmitted")
        report_tag = "datesubmitted" if "datesubmitted" in row_sdts else ("datesubmited" if "datesubmited" in row_sdts else None)
        if dr_col and report_tag:
            _set_sdt_text(row_sdts[report_tag], format_date_like(record.get(dr_col, "")))

def build_output_filename(bill_anchor: date, form_no: int) -> str:
    return FILENAME_PATTERN.format(
        BILL_MONTH_NUM=bill_anchor.month,
        BILL_MONTH_NAME_SHORT=month_short_name(bill_anchor),
        BILL_MONTH_NAME_LONG=bill_anchor.strftime("%B"),
        BILL_YEAR=bill_anchor.year,
        FORM_NO=form_no,
    )

def main():
    # Choose the month to compile
    bill_start, use_mtd = prompt_for_month()

    print(f"Loading Excel (read-only) from {EXCEL_PATH} …")
    df = load_rows_for_window(bill_start, use_mtd)
    df = load_rows_for_window(bill_start, use_mtd)

    resolved_map: Dict[int, str] = df.attrs["resolved_map"]

    if df.empty:
        print("No rows found for the chosen period. Nothing to do.")
        return

    # Prepare output
    ensure_output_dir(OUTPUT_DIR)

    # Records and pagination (7 per form)
    records: List[Dict[str, Any]] = df.to_dict(orient="records")
    pages = [records[i:i+7] for i in range(0, len(records), 7)]

    for idx, page_records in enumerate(pages, start=1):
        print(f"Building form {idx} with {len(page_records)} visit(s)…")
        doc = Document(TEMPLATE_PATH)

        # Auto-locate the correct table (header-based, min 8 rows)
        table_idx = find_target_table(doc)

        # Fill the table rows (no structural changes)
        fill_table(doc, page_records, resolved_map, table_idx)

        # Save
        out_name = build_output_filename(bill_start, idx)
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        print(f"  Saved: {out_path}")

    print("Done.")

if __name__ == "__main__":
    main()
