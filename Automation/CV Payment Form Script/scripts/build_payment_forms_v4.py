"""
Court Visitor Payment form auto-filler (v4)
- Reads Excel at EXCEL_PATH (read-only; never writes back)
- Prompts for month (accepts many formats; see parse_month_input)
- Filters to that month (optionally month-to-date)
- Auto-selects the correct template table (requires >= 8 rows and header match; prefers table containing 'C-1-PB-')
- Does NOT change the template structure (no adding/removing rows)
- Handles the split 'Cause No.' column: a fixed 'C-1-PB-' cell + a separate cell for the suffix
- Saves outputs as "9_1Sept2025 Court Visitor Payment.docx", etc., in OUTPUT_DIR

Install once:
  python -m pip install pandas python-docx python-dateutil openpyxl
"""
from __future__ import annotations
import os
import re
from dataclasses import dataclass
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
from typing import List, Dict, Any, Optional, Sequence

import pandas as pd
from docx import Document

# ===========================
# CONFIG — EDIT THESE IF NEEDED
# ===========================
EXCEL_PATH = r"C:\Users\may\OneDrive\Guardian Docs\ward_guardian_info.xlsx"  # READ-ONLY
SHEET_NAME = 0  # first sheet
DATE_COLUMN = "visitdate"  # used to select the month
OUTPUT_DIR = r"G:\My Drive\Guardianship files\Payments to review and submit"
TEMPLATE_PATH = r"C:\GuardianAutomation\2025 Court Visitor Payment Invoice - fields.docx"  # your local .docx

# Header bits to identify the right table (case-insensitive contains)
EXPECTED_HEADER_BITS = [
    "cause",
    "date appointed",
    "date of court visit",
    "date of court visitor report",
]

# Map (logical field -> Excel aliases). We'll place these at detected column offsets.
EXCEL_ALIASES: Dict[str, Sequence[str]] = {
    "cause": ("causeno",),
    "date_appointed": ("Dateappointed", "dateappointed"),
    "date_visit": ("visitdate",),
    "date_report": ("datesubmitted", "datesubmited"),
}

# Date formatting on the form
DATE_FORMAT = "%m/%d/%Y"  # e.g., 09/05/2025

# Output file name pattern
FILENAME_PATTERN = "{BILL_MONTH_NUM}_{FORM_NO}{BILL_MONTH_NAME_SHORT}{BILL_YEAR} Court Visitor Payment.docx"

# ===========================
# END CONFIG
# ===========================

def month_short_name(dt: date) -> str:
    """Return short month name with 'Sept' for September (instead of 'Sep')."""
    s = dt.strftime("%b")
    return "Sept" if s == "Sep" else s

def parse_month_input(raw: str) -> Optional[date]:
    """
    Accepts: 'last', '', 'YYYY-MM', 'YYYY/MM', 'MM-YYYY', 'MM/YYYY', 'M-YYYY', 'M/YYYY',
             'YYYY-M', 'YYYY/M', '9', '09'
    Returns date set to the FIRST of the target month, or None if unparseable.
    """
    raw = (raw or "").strip().lower()
    today = date.today()

    if raw == "":
        return today.replace(day=1)
    if raw == "last":
        return (today.replace(day=1) - relativedelta(months=1))

    if raw.isdigit():
        m = int(raw)
        if 1 <= m <= 12:
            return date(today.year, m, 1)

    # Normalize things like '2025-9', '9-2025', '2025/9', '09/2025'
    m = re.fullmatch(r"(\d{4})[-/](\d{1,2})", raw)
    if m:
        y, mon = int(m.group(1)), int(m.group(2))
        if 1 <= mon <= 12:
            return date(y, mon, 1)
    m = re.fullmatch(r"(\d{1,2})[-/](\d{4})", raw)
    if m:
        mon, y = int(m.group(1)), int(m.group(2))
        if 1 <= mon <= 12:
            return date(y, mon, 1)

    for fmt in ("%Y-%m", "%Y/%m", "%m-%Y", "%m/%Y"):
        try:
            dt = datetime.strptime(raw, fmt).date()
            return dt.replace(day=1)
        except ValueError:
            pass

    return None

def prompt_for_month() -> (date, bool):
    raw = input(
        "Enter month to compile (e.g., 9/2025, 09-2025, 2025-9, 2025/9, YYYY-MM). "
        "Blank = current month, or 'last' = previous month: "
    )
    anchor = parse_month_input(raw)
    if anchor is None:
        print("Could not parse; defaulting to current month.")
        anchor = date.today().replace(day=1)

    today = date.today()
    use_mtd = False
    if anchor.year == today.year and anchor.month == today.month:
        yn = input("Use month-to-date for the current month? [Y/n]: ").strip().lower()
        use_mtd = (yn in ("", "y", "yes"))
    else:
        yn = input("Compile the FULL month (not capped to today)? [Y/n]: ").strip().lower()
        use_mtd = not (yn in ("", "y", "yes"))
    return anchor, use_mtd

@dataclass
class BillingWindow:
    start: date
    end: date  # exclusive

    @staticmethod
    def for_month(anchor: date, to_date: Optional[date] = None) -> "BillingWindow":
        start = anchor.replace(day=1)
        end = (start + relativedelta(months=1))
        if to_date is not None:
            end = min(end, to_date + timedelta(days=1))
        return BillingWindow(start, end)

def ensure_output_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def first_existing_col(df: pd.DataFrame, candidates: Sequence[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    lower_map = {c.lower(): c for c in df.columns}
    for c in candidates:
        if c.lower() in lower_map:
            return lower_map[c.lower()]
    return None

def load_rows_for_window(bill_start: date, use_month_to_date: bool) -> pd.DataFrame:
    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, engine="openpyxl")

    if DATE_COLUMN not in df.columns:
        raise ValueError(f"Expected date column '{DATE_COLUMN}' not found in Excel. Found: {list(df.columns)}")

    df[DATE_COLUMN] = pd.to_datetime(df[DATE_COLUMN], errors="coerce").dt.date

    to_date_limit = date.today() if use_month_to_date else None
    window = BillingWindow.for_month(bill_start, to_date=to_date_limit)
    mask = (df[DATE_COLUMN] >= window.start) & (df[DATE_COLUMN] < window.end)
    month_df = df.loc[mask].copy()

    sort_keys = [DATE_COLUMN]
    if "causeno" in month_df.columns:
        sort_keys.append("causeno")
    month_df.sort_values(sort_keys, inplace=True)

    # Resolve Excel headers for each logical field
    resolved: Dict[str, str] = {}
    for logical, aliases in EXCEL_ALIASES.items():
        col = first_existing_col(month_df, aliases)
        if not col:
            raise ValueError(
                f"None of the aliases {aliases} exist in the Excel header. "
                f"Headers present: {list(month_df.columns)}"
            )
        resolved[logical] = col
    month_df.attrs["resolved"] = resolved

    return month_df.reset_index(drop=True)

def find_target_table(doc: Document) -> int:
    """
    Choose the table that:
      - has >= 8 rows (1 header + 7 data rows),
      - has >= 4 columns,
      - header contains EXPECTED_HEADER_BITS (case-insensitive),
      - and preferably contains 'C-1-PB' in any data row.
    """
    best_idx = None
    best_score = -1

    for idx, t in enumerate(doc.tables):
        row_count = len(t.rows)
        col_count = len(t.columns) if row_count else 0
        if row_count < 8 or col_count < 4:
            continue

        try:
            header_text = " | ".join(cell.text.strip() for cell in t.rows[0].cells).lower()
        except Exception:
            header_text = ""

        score = sum(bit in header_text for bit in EXPECTED_HEADER_BITS)

        # Presence of 'C-1-PB' in any of the first few data rows is a strong hint
        has_prefix = False
        for r in range(1, min(row_count, 8)):
            for cell in t.rows[r].cells:
                if "c-1-pb" in cell.text.lower():
                    has_prefix = True
                    break
            if has_prefix:
                break
        if has_prefix:
            score += 10  # big boost

        if score > best_score:
            best_score = score
            best_idx = idx

    return best_idx if best_idx is not None else 0

def format_cell_value(field_name: str, val: Any) -> str:
    if pd.isna(val):
        return ""
    # Date fields
    if field_name in ("date_appointed", "date_visit", "date_report"):
        try:
            if isinstance(val, date):
                return val.strftime(DATE_FORMAT)
            if isinstance(val, datetime):
                return val.date().strftime(DATE_FORMAT)
            parsed = pd.to_datetime(val, errors="coerce")
            if pd.notna(parsed):
                return parsed.date().strftime(DATE_FORMAT)
        except Exception:
            pass
    return str(val)

def cause_suffix_from_value(raw: Any) -> str:
    """
    Return the portion of the cause number AFTER the 'C-1-PB-' prefix.
    If the Excel value already includes the prefix, strip it.
    """
    if pd.isna(raw):
        return ""
    s = str(raw).strip()
    # Normalize and strip any leading prefix variants like 'C-1-PB-' or 'C-1-PB'
    s_norm = re.sub(r"^\s*(?i:c-?1-?pb-?)\s*", "", s)
    return s_norm

def clear_table_data_area(t) -> None:
    """Clear rows 1..7 (keep header at row 0) without changing structure."""
    max_data_rows = min(len(t.rows) - 1, 7)
    for r in range(1, 1 + max_data_rows):
        for c in range(len(t.columns)):
            try:
                cell = t.cell(r, c)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                if not cell.paragraphs:
                    cell.add_paragraph("")
            except Exception:
                pass

def detect_column_positions(t) -> Optional[Dict[str, int]]:
    """
    Detect actual column indices:
      - Find the cell containing 'C-1-PB' in the FIRST data row (row 1).
      - Then assume cause suffix is the NEXT column,
        and the three dates follow in order.
    Returns a dict with keys: 'cause', 'date_appointed', 'date_visit', 'date_report'.
    """
    if len(t.rows) < 2:
        return None

    # Search row 1 for the 'C-1-PB' prefix cell
    prefix_col = None
    for c_idx in range(len(t.columns)):
        try:
            txt = t.cell(1, c_idx).text.strip().lower()
        except Exception:
            txt = ""
        if "c-1-pb" in txt:
            prefix_col = c_idx
            break

    if prefix_col is None:
        return None  # fallback handled by caller

    # Cause suffix is next column; then three date columns
    return {
        "cause": prefix_col + 1,
        "date_appointed": prefix_col + 2,
        "date_visit": prefix_col + 3,
        "date_report": prefix_col + 4,
    }

def fill_table(doc: Document, records: List[Dict[str, Any]], resolved_excel: Dict[str, str], table_index: int) -> None:
    tables = doc.tables
    if not tables:
        raise RuntimeError("Template has no tables.")

    try:
        t = tables[table_index]
    except IndexError:
        raise RuntimeError(f"Template does not have table index {table_index}; found {len(tables)} tables.")

    required_rows = 8  # header + 7 data rows
    if len(t.rows) < required_rows:
        raise RuntimeError(
            f"Selected table only has {len(t.rows)} rows; need at least {required_rows} "
            f"(header + 7 lines)."
        )

    # Determine actual column positions based on the 'C-1-PB-' prefix cell
    col_pos = detect_column_positions(t)
    if not col_pos:
        # Fallback: try to infer from the header row (less reliable on merged headers)
        header_map = {}
        try:
            headers = [cell.text.strip().lower() for cell in t.rows[0].cells]
        except Exception:
            headers = []
        for idx, text in enumerate(headers):
            if "cause" in text and "visitor" not in text:
                header_map["cause"] = idx
            elif "date appointed" in text:
                header_map["date_appointed"] = idx
            elif "date of court visit" in text:
                header_map["date_visit"] = idx
            elif "date of court visitor report" in text:
                header_map["date_report"] = idx
        # Only use this fallback if we got all four
        if all(k in header_map for k in ("cause", "date_appointed", "date_visit", "date_report")):
            col_pos = header_map
        else:
            raise RuntimeError("Could not detect the correct columns in the table. Check the template structure.")

    # Clear existing data
    clear_table_data_area(t)

    # Fill up to 7 rows
    for row_idx, record in enumerate(records[:7], start=1):
        # Cause suffix
        cause_val = record.get(resolved_excel["cause"], "")
        suffix = cause_suffix_from_value(cause_val)
        try:
            cell = t.cell(row_idx, col_pos["cause"])
            if cell.paragraphs:
                cell.paragraphs[0].text = suffix
            else:
                cell.add_paragraph(suffix)
        except Exception:
            pass

        # Dates
        for key in ("date_appointed", "date_visit", "date_report"):
            excel_col = resolved_excel[key]
            val = record.get(excel_col, "")
            txt = format_cell_value(key, val)
            try:
                cell = t.cell(row_idx, col_pos[key])
                if cell.paragraphs:
                    cell.paragraphs[0].text = txt
                else:
                    cell.add_paragraph(txt)
            except Exception:
                pass

def build_output_filename(bill_anchor: date, form_no: int) -> str:
    return FILENAME_PATTERN.format(
        BILL_MONTH_NUM=bill_anchor.month,
        BILL_MONTH_NAME_SHORT=month_short_name(bill_anchor),
        BILL_MONTH_NAME_LONG=bill_anchor.strftime("%B"),
        BILL_YEAR=bill_anchor.year,
        FORM_NO=form_no,
    )

def main():
    # Choose the month to compile
    bill_start, use_mtd = prompt_for_month()

    print(f"Loading Excel (read-only) from {EXCEL_PATH} …")
    df = load_rows_for_window(bill_start, use_mtd)
    if df.empty:
        print("No rows found for the chosen period. Nothing to do.")
        return

    resolved_excel: Dict[str, str] = df.attrs["resolved"]

    # Prepare output
    ensure_output_dir(OUTPUT_DIR)

    # Records and pagination (7 per form)
    records: List[Dict[str, Any]] = df.to_dict(orient="records")
    pages = [records[i:i+7] for i in range(0, len(records), 7)]

    for idx, page_records in enumerate(pages, start=1):
        print(f"Building form {idx} with {len(page_records)} visit(s)…")
        doc = Document(TEMPLATE_PATH)

        # Auto-locate the correct table
        table_idx = find_target_table(doc)

        # Fill (no structural changes)
        fill_table(doc, page_records, resolved_excel, table_idx)

        # Save
        out_name = build_output_filename(bill_start, idx)
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        print(f"  Saved: {out_path}")

    print("Done.")

if __name__ == "__main__":
    main()
