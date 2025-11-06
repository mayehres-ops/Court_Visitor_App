"""
Court Visitor Payment form auto-filler (openpyxl + SDT, date-control safe)
Paths updated for GoogleSync layout.

Excel:   C:\GoogleSync\GuardianShip_App\App Data\ward_guardian_info.xlsx
Template:C:\GoogleSync\GuardianShip_App\Templates\Court_Visitor_Payment_Invoice.docx
Output:  C:\GoogleSync\GuardianShip_App\App Data\Output\Payment Forms
Script:  C:\GoogleSync\GuardianShip_App\Automation\CV Payment Form Script\scripts\build_payment_forms_sdt.py

Install once:
  python -m pip install python-docx openpyxl python-dateutil
"""

import os
import re
import sys
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
from typing import List, Dict, Any, Optional, Sequence
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.oxml import OxmlElement

# ===========================
# CONFIG — YOUR NEW LOCATIONS
# ===========================

# Dynamic path detection
_script_dir = Path(__file__).parent.parent.parent.parent  # Go up to app root
try:
    sys.path.insert(0, str(_script_dir / "Scripts"))
    from app_paths import get_app_paths
    from cv_info_manager import get_cv_info
    _app_paths = get_app_paths(str(_script_dir))

    EXCEL_PATH = str(_app_paths.EXCEL_PATH)
    TEMPLATE_PATH = str(_app_paths.APP_ROOT / "Templates" / "Court_Visitor_Payment_Invoice.docx")
    OUTPUT_DIR = str(_app_paths.APP_ROOT / "App Data" / "Output" / "Payment Forms")

except Exception:
    # Fallback to hardcoded paths
    EXCEL_PATH = r"C:\GoogleSync\GuardianShip_App\App Data\ward_guardian_info.xlsx"
    TEMPLATE_PATH = r"C:\GoogleSync\GuardianShip_App\Templates\Court_Visitor_Payment_Invoice.docx"
    OUTPUT_DIR = r"C:\GoogleSync\GuardianShip_App\App Data\Output\Payment Forms"

# Header bits we look for in the table header row (case-insensitive)
EXPECTED_HEADER_BITS = [
    "cause",
    "date appointed",
    "date of court visit",
    "date of court visitor report",
]

# Excel header aliases (we normalize and resolve to whatever you actually have)
VISIT_KEYS  = ("visitdate", "visit date", "date of court visit", "court visit date")
CAUSE_KEYS  = ("causeno", "cause no", "cause", "case no", "caseno")
APPT_KEYS   = ("dateappointed", "date appointed", "appointment date", "appointmentdate", "date of appointment")
SUBMIT_KEYS = ("datesubmitted", "datesubmited", "date submitted", "date of court visitor report",
               "visitor report date", "report date", "reportdate", "date of visitor report")

DATE_FORMAT = "%m/%d/%Y"
FILENAME_PATTERN = "{BILL_MONTH_NUM}_{FORM_NO}{BILL_MONTH_NAME_SHORT}{BILL_YEAR} Court Visitor Payment.docx"

# OOXML namespace (Clark notation)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"

# Toggle extra prints
DEBUG = True
# ===========================


def month_short_name(dt):
    s = dt.strftime("%b")
    return "Sept" if s == "Sep" else s


def parse_month_input(raw):
    raw = (raw or "").strip().lower()
    today = date.today()
    if raw == "":
        return today.replace(day=1)
    if raw == "last":
        return (today.replace(day=1) - relativedelta(months=1))
    if raw.isdigit():
        m = int(raw)
        if 1 <= m <= 12:
            return date(today.year, m, 1)
    m = re.fullmatch(r"(\d{4})[-/](\d{1,2})", raw)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12:
            return date(y, mo, 1)
    m = re.fullmatch(r"(\d{1,2})[-/](\d{4})", raw)
    if m:
        mo, y = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12:
            return date(y, mo, 1)
    for fmt in ("%Y-%m", "%Y/%m", "%m-%Y", "%m/%Y"):
        try:
            return datetime.strptime(raw, fmt).date().replace(day=1)
        except ValueError:
            pass
    return None


def prompt_for_month():
    """Show GUI picker to select month for payment forms"""

    class MonthPickerDialog:
        def __init__(self):
            self.result = None
            self.root = tk.Tk()
            self.root.title("Select Month for Payment Forms")
            self.root.geometry("450x320")
            self.root.resizable(False, False)

            # Center the window
            self.root.update_idletasks()
            x = (self.root.winfo_screenwidth() // 2) - (450 // 2)
            y = (self.root.winfo_screenheight() // 2) - (320 // 2)
            self.root.geometry(f'+{x}+{y}')

            # Title
            title_label = tk.Label(self.root, text="Generate Payment Forms",
                                 font=('Segoe UI', 14, 'bold'))
            title_label.pack(pady=15)

            # Month selection frame
            frame = ttk.Frame(self.root, padding=10)
            frame.pack(fill='both', expand=True)

            # Month dropdown
            ttk.Label(frame, text="Select Month:", font=('Segoe UI', 10)).grid(row=0, column=0, sticky='w', pady=5)

            # Generate month options (current month + 11 previous months)
            # The picker automatically updates each month (e.g., in November, "current" will be November)
            today = date.today()
            self.month_options = []
            for i in range(12):  # Start from 0 (current month) to 11 (12 months ago)
                month_date = today.replace(day=1) + relativedelta(months=-i)
                label = month_date.strftime("%B %Y")
                if i == 0:
                    label += " (current)"
                elif i == 1:
                    label += " (previous)"
                self.month_options.append((label, month_date))

            self.month_var = tk.StringVar(value=self.month_options[0][0])  # Default to current month
            month_dropdown = ttk.Combobox(frame, textvariable=self.month_var,
                                        values=[m[0] for m in self.month_options],
                                        state='readonly', width=35)
            month_dropdown.grid(row=0, column=1, pady=5, padx=10)

            # Month type (full or month-to-date)
            ttk.Label(frame, text="Period:", font=('Segoe UI', 10)).grid(row=1, column=0, sticky='w', pady=5)

            self.period_var = tk.StringVar(value="full")
            ttk.Radiobutton(frame, text="Full month", variable=self.period_var,
                          value="full").grid(row=1, column=1, sticky='w', padx=10)
            ttk.Radiobutton(frame, text="Month-to-date", variable=self.period_var,
                          value="mtd").grid(row=2, column=1, sticky='w', padx=10)

            # Info label
            info_label = ttk.Label(frame, text="Payment forms are typically generated the following month\n"
                                             "for the previous month's completed work.",
                                 font=('Segoe UI', 9), foreground='gray')
            info_label.grid(row=3, column=0, columnspan=2, pady=15)

            # Buttons
            btn_frame = ttk.Frame(self.root)
            btn_frame.pack(pady=10)

            ttk.Button(btn_frame, text="Generate", command=self.on_generate,
                      width=12).pack(side='left', padx=5)
            ttk.Button(btn_frame, text="Cancel", command=self.on_cancel,
                      width=12).pack(side='left', padx=5)

            self.root.protocol("WM_DELETE_WINDOW", self.on_cancel)

        def on_generate(self):
            selected_label = self.month_var.get()
            # Find the matching date
            for label, month_date in self.month_options:
                if label == selected_label:
                    use_mtd = (self.period_var.get() == "mtd")
                    self.result = (month_date, use_mtd)
                    break
            self.root.destroy()

        def on_cancel(self):
            self.result = None
            self.root.destroy()

        def show(self):
            self.root.mainloop()
            return self.result

    # Show picker dialog
    picker = MonthPickerDialog()
    result = picker.show()

    if result is None:
        print("Payment form generation cancelled by user")
        sys.exit(0)

    anchor, use_mtd = result
    period_type = "month-to-date" if use_mtd else "full month"
    print(f"Selected: {anchor.strftime('%B %Y')} ({period_type})")
    return anchor, use_mtd


# -------- Excel helpers (openpyxl) --------
def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(s or "").strip().lower())

def _find_header_map(headers: List[str]) -> Dict[str, str]:
    """Return mapping of logical names to actual header strings in the sheet."""
    norm_to_real = {_norm(h): h for h in headers}
    def first_match(keys: Sequence[str]) -> Optional[str]:
        for k in keys:
            nk = _norm(k)
            if nk in norm_to_real:
                return norm_to_real[nk]
        return None

    visit = first_match(VISIT_KEYS)
    cause = first_match(CAUSE_KEYS)
    appt  = first_match(APPT_KEYS)
    submit = first_match(("datesubmitted",) + SUBMIT_KEYS)  # prefer the corrected spelling

    missing = [n for n, v in [("visitdate", visit), ("cause", cause), ("dateappointed", appt), ("datesubmitted", submit)] if not v]
    if missing:
        raise ValueError(f"Missing expected columns: {missing}. Headers present: {headers}")

    return {"visitdate": visit, "cause": cause, "dateappointed": appt, "datesubmitted": submit}

def _to_date(val) -> Optional[date]:
    """Parse any cell value to a date (supports python date/datetime and common strings)."""
    if val is None:
        return None
    if isinstance(val, date) and not isinstance(val, datetime):
        return val
    if isinstance(val, datetime):
        return val.date()
    s = str(val).strip()
    if not s:
        return None
    # Try common patterns
    for fmt in ("%m/%d/%Y", "%m-%d-%Y", "%Y-%m-%d", "%Y/%m/%d", "%m/%d/%y", "%m-%d-%y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    # Extract first mm/dd/yyyy pattern if the cell has extra text (e.g., "9/10/2025 (submitted)")
    m = re.search(r"\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b", s)
    if m:
        mon, day, yr = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if yr < 100:  # normalize 2-digit year
            yr += 2000 if yr < 50 else 1900
        try:
            return date(yr, mon, day)
        except Exception:
            pass
    # ISO-ish fallback
    try:
        return datetime.fromisoformat(s).date()
    except Exception:
        return None

def read_sheet_rows(path: str, sheet_index: int = 0) -> (List[str], List[Dict[str, Any]]):
    """Read entire sheet as list of dicts with real header names using openpyxl (data_only)."""
    wb = load_workbook(path, data_only=True)
    ws = wb.worksheets[sheet_index]
    headers = [c.value if c.value is not None else "" for c in next(ws.iter_rows(min_row=1, max_row=1))]
    rows: List[Dict[str, Any]] = []
    for r in ws.iter_rows(min_row=2, values_only=True):
        row = {}
        for h, v in zip(headers, r):
            row[h] = v
        rows.append(row)
    return headers, rows


def load_rows_for_window(bill_start: date, use_month_to_date: bool):
    headers, rows = read_sheet_rows(EXCEL_PATH, 0)
    header_map = _find_header_map(headers)

    # Filter by visitdate month
    start = bill_start.replace(day=1)
    end   = (start + relativedelta(months=1))
    if use_month_to_date:
        end = min(end, date.today() + timedelta(days=1))

    filtered: List[Dict[str, Any]] = []
    for row in rows:
        v = _to_date(row.get(header_map["visitdate"]))
        if v is None:
            continue
        if start <= v < end:
            filtered.append(row)

    # Sort by visitdate then cause
    def sort_key(r):
        return (_to_date(r.get(header_map["visitdate"])) or date.min, str(r.get(header_map["cause"]) or ""))

    filtered.sort(key=sort_key)

    if DEBUG:
        print("Resolved Excel headers:", header_map)
        print("Found", len(filtered), "visit row(s) matching the period.")

    return header_map, filtered


# ---------- Word: find table & write SDTs ----------
def find_target_table(doc: Document) -> int:
    best_idx, best_score = None, -1
    fallback_idx, fallback_rows = None, -1
    for idx, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.columns) if rows else 0
        if rows > fallback_rows:
            fallback_rows, fallback_idx = rows, idx
        if rows < 8:  # header + 7 data rows
            continue
        try:
            header_text = " | ".join(c.text.strip().lower() for c in t.rows[0].cells)
        except Exception:
            header_text = ""
        score = sum(bit in header_text for bit in EXPECTED_HEADER_BITS)
        if score > best_score:
            best_score, best_idx = score, idx
    return best_idx if best_idx is not None else (fallback_idx if fallback_idx is not None else 0)


def _iter_sdts(element):
    for el in element.iter():
        if isinstance(el.tag, str) and el.tag.endswith("}sdt"):
            yield el

def _set_sdt_text(sdt_elm, new_text):
    if new_text is None:
        new_text = ""
    sdt_content = sdt_elm.find(".//" + w("sdtContent"))
    if sdt_content is None:
        return
    texts = sdt_content.findall(".//" + w("t"))
    if texts:
        texts[-1].text = new_text
        for t in texts[:-1]:
            t.text = ""
        return
    p = OxmlElement(w("p"))
    r = OxmlElement(w("r"))
    t = OxmlElement(w("t"))
    t.text = new_text
    r.append(t)
    p.append(r)
    sdt_content.append(p)

def _set_sdt_date_value(sdt_elm, dt_obj: Optional[date]):
    """If SDT is a DATE control, set w:fullDate; also sets visible text."""
    if dt_obj is None:
        _set_sdt_text(sdt_elm, "")
        return
    _set_sdt_text(sdt_elm, dt_obj.strftime(DATE_FORMAT))
    sdt_pr = sdt_elm.find(".//" + w("sdtPr"))
    if sdt_pr is None:
        return
    date_el = sdt_pr.find(".//" + w("date"))
    if date_el is None:
        return
    full = date_el.find(".//" + w("fullDate"))
    if full is None:
        full = OxmlElement(w("fullDate"))
        date_el.append(full)
    full.set(w("val"), dt_obj.strftime("%Y-%m-%dT00:00:00Z"))

def _cause_suffix(raw) -> str:
    s = str(raw or "").strip()
    return re.sub(r"^\s*(?i:c-?1-?pb-?)\s*", "", s)

def fill_cv_info(doc: Document):
    """Fill Court Visitor information from config into content controls."""
    try:
        cv_info = get_cv_info()

        # Map config fields to content control tags
        # User should name their content controls accordingly in the Word template
        tag_mapping = {
            'Court Visitor Name': cv_info.get('name', ''),
            'CV Name': cv_info.get('name', ''),
            'Vendor Number': cv_info.get('vendor_number', ''),
            'Vendor No': cv_info.get('vendor_number', ''),
            'Address Line 1': cv_info.get('address_line1', ''),
            'Address 1': cv_info.get('address_line1', ''),
            'Address Line 2': cv_info.get('address_line2', ''),
            'Address 2': cv_info.get('address_line2', ''),
        }

        # Iterate through all SDTs in the document and fill matching ones
        for sdt_elm in _iter_sdts(doc.element.body):
            # Try to get the tag/title from sdtPr
            sdt_pr = sdt_elm.find(".//" + w("sdtPr"))
            if sdt_pr is None:
                continue

            # Check for tag element
            tag_elem = sdt_pr.find(".//" + w("tag"))
            if tag_elem is not None:
                tag_val = tag_elem.get(w("val"), "")
                if tag_val in tag_mapping:
                    _set_sdt_text(sdt_elm, tag_mapping[tag_val])
                    if DEBUG:
                        print(f"  Filled CV content control '{tag_val}': {tag_mapping[tag_val]}")
                    continue

            # Check for alias (title) element
            alias_elem = sdt_pr.find(".//" + w("alias"))
            if alias_elem is not None:
                alias_val = alias_elem.get(w("val"), "")
                if alias_val in tag_mapping:
                    _set_sdt_text(sdt_elm, tag_mapping[alias_val])
                    if DEBUG:
                        print(f"  Filled CV content control '{alias_val}': {tag_mapping[alias_val]}")

        if DEBUG:
            print("Court Visitor info filled successfully")

    except Exception as e:
        # Gracefully continue if CV config not available
        print(f"Note: Could not load CV info: {e}")

def fill_table(doc: Document, table_idx: int, header_map: Dict[str, str], page_rows: List[Dict[str, Any]]):
    t = doc.tables[table_idx]
    for row_idx, row in enumerate(page_rows[:7], start=1):
        tr = t.rows[row_idx]._tr
        sdts = list(_iter_sdts(tr))
        if DEBUG:
            cause_val = row.get(header_map["cause"])
            appt_val  = row.get(header_map["dateappointed"])
            visit_val = row.get(header_map["visitdate"])
            subm_val  = row.get(header_map["datesubmitted"])
            print(f"  Row {row_idx}: cause='{cause_val}'  dateappointed='{appt_val}'  visitdate='{visit_val}'  report='{subm_val}'")

        if len(sdts) >= 1:
            _set_sdt_text(sdts[0], _cause_suffix(row.get(header_map["cause"])))
        if len(sdts) >= 2:
            _set_sdt_date_value(sdts[1], _to_date(row.get(header_map["dateappointed"])))
        if len(sdts) >= 3:
            _set_sdt_date_value(sdts[2], _to_date(row.get(header_map["visitdate"])))
        if len(sdts) >= 4:
            _set_sdt_date_value(sdts[3], _to_date(row.get(header_map["datesubmitted"])))


def build_output_filename(bill_anchor: date, form_no: int) -> str:
    return FILENAME_PATTERN.format(
        BILL_MONTH_NUM=bill_anchor.month,
        BILL_MONTH_NAME_SHORT=month_short_name(bill_anchor),
        BILL_YEAR=bill_anchor.year,
        FORM_NO=form_no,
    )


def main():
    print(">>> Running script:", __file__)
    # quick sanity checks so path mistakes are obvious
    if not os.path.isfile(EXCEL_PATH):
        print(f"[ERROR] Excel not found at:\n  {EXCEL_PATH}")
        sys.exit(1)
    if not os.path.isfile(TEMPLATE_PATH):
        print(f"[ERROR] Template not found at:\n  {TEMPLATE_PATH}")
        sys.exit(1)

    bill_start, use_mtd = prompt_for_month()

    print(f"Loading Excel (raw, read-only) from {EXCEL_PATH} …")
    header_map, rows = load_rows_for_window(bill_start, use_mtd)
    print("Resolved Excel headers:", header_map)
    print(f"Found {len(rows)} visit row(s) matching the period.")

    if not rows:
        print("[WARN] No rows found for the chosen period. Nothing to do.")
        print("\n[OK] Payment Forms - No forms to generate (no visits in selected period)")
        sys.exit(0)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Chunk into pages of 7
    pages = [rows[i:i+7] for i in range(0, len(rows), 7)]

    generated_files = []
    for idx, page_rows in enumerate(pages, start=1):
        print(f"Building form {idx} with {len(page_rows)} visit(s)…")
        doc = Document(TEMPLATE_PATH)
        table_idx = find_target_table(doc)
        if DEBUG:
            print("Using table index:", table_idx)

        # Fill Court Visitor information
        fill_cv_info(doc)

        # Fill table with visit data
        fill_table(doc, table_idx, header_map, page_rows)
        out_name = build_output_filename(bill_start, idx)
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        generated_files.append(out_path)
        print(f"  Saved: {out_path}")

    print("Done.")

    # Open the generated files for printing
    if generated_files:
        print(f"\nOpening {len(generated_files)} payment form(s)...")
        import subprocess
        for file_path in generated_files:
            try:
                # Use subprocess with shell=True for better compatibility when called from GUI
                subprocess.Popen(['start', '', file_path], shell=True)
                print(f"  Opened: {file_path}")
            except Exception as e:
                # Fallback to os.startfile
                try:
                    os.startfile(file_path)
                    print(f"  Opened: {file_path}")
                except Exception as e2:
                    print(f"  WARNING: Could not open {file_path}: {e} / {e2}")

    # Exit with proper success code
    print(f"\n[OK] Payment Forms SUCCESS - Generated {len(generated_files)} form(s)")
    sys.exit(0)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n[FAIL] Payment Forms FAILED: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
