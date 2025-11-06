"""
Court Visitor Payment form auto-filler (v2)

What this script does
- Opens your Excel READ-ONLY at EXCEL_PATH (never writes back)
- Prompts you for which month to compile (supports full month or month‑to‑date)
- Filters rows by that month using your visitdate column
- Fills the Word template table with 7 lines per form, left→right columns:
    Cause No. | Date Appointed | Date of Court Visit | Date of Court Visitor Report
- Creates as many forms as needed (7 visits per form)
- Saves files like "9_1Sept2025 Court Visitor Payment.docx" into
  OUTPUT_DIR

Setup quick-start
1) Install dependencies once:
   pip install pandas python-docx python-dateutil openpyxl

2) Update CONFIG below if your local paths differ.

3) Double-click your .bat launcher to run (see example at bottom of this file).

Safe by design
- Excel is opened READ-ONLY via pandas; this script never modifies it.
- All output is new .docx files in OUTPUT_DIR.

Change log (v2)
- Maps EXACT columns you specified
- Robust to your Excel typos/variants (datesubmitted vs datesubmited)
- Prompts for target month (YYYY-MM, blank = current, "last" = previous)
- Option to cap to month‑to‑date for current month
- Auto-detects the correct Word table by matching header text; falls back to TABLE_INDEX
"""
from __future__ import annotations
import os
from dataclasses import dataclass
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
from typing import List, Dict, Any, Optional, Sequence

import pandas as pd
from docx import Document

# ===========================
# CONFIG — EDIT THESE IF NEEDED
# ===========================
EXCEL_PATH = r"C:\\Users\\may\\OneDrive\\Guardian Docs\\ward_guardian_info.xlsx"  # DO NOT CHANGE IT PROGRAMMATICALLY
SHEET_NAME = 0  # first sheet
DATE_COLUMN = "visitdate"  # used to select the month
OUTPUT_DIR = r"G:\\My Drive\\Guardianship files\\Payments to review and submit"
TEMPLATE_PATH = r"C:\\GuardianAutomation\\2025 Court Visitor Payment Invoice - fields.docx"  # put your local copy here

# If auto-detect fails, we fall back to this table index (0 = first)
TABLE_INDEX = 0

# WORD TABLE HEADERS we expect (case-insensitive contains match). These help locate the correct table.
EXPECTED_HEADERS = [
    "Cause",            # Cause No.
    "Date Appointed",  # Date Appointed
    "Date of Court Visit",
    "Date of Court Visitor Report",
]

# The 4 table columns mapping (left→right) to Excel headers. We’ll resolve aliases below.
COLUMN_MAP_BASE: Dict[int, Sequence[str]] = {
    0: ("causeno",),              # Cause No.
    1: ("Dateappointed", "dateappointed"),  # Date Appointed
    2: ("visitdate",),            # Date of Court Visit
    3: ("datesubmitted", "datesubmited"),  # Date of Court Visitor Report
}

# Date formatting on the form
DATE_FORMAT = "%m/%d/%Y"  # e.g., 09/05/2025

# Output file name pattern
FILENAME_PATTERN = "{BILL_MONTH_NUM}_{FORM_NO}{BILL_MONTH_NAME_SHORT}{BILL_YEAR} Court Visitor Payment.docx"

# ===========================
# END CONFIG
# ===========================

@dataclass
class BillingWindow:
    start: date
    end: date  # exclusive

    @staticmethod
    def for_month(anchor: date, to_date: Optional[date] = None) -> "BillingWindow":
        start = anchor.replace(day=1)
        end = (start + relativedelta(months=1))
        if to_date is not None:
            # Cap end at the day AFTER to_date to keep half-open interval logic
            end = min(end, to_date + timedelta(days=1))
        return BillingWindow(start, end)


def ensure_output_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def first_existing_col(df: pd.DataFrame, candidates: Sequence[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    # try case-insensitive match
    lower_map = {c.lower(): c for c in df.columns}
    for c in candidates:
        if c.lower() in lower_map:
            return lower_map[c.lower()]
    return None


def prompt_for_month() -> (date, bool):
    """Ask user which month to compile. Returns (anchor_date_first_of_month, use_month_to_date)."""
    raw = input("Enter month to compile (YYYY-MM), blank = current month, or 'last' = previous month: ").strip()
    today = date.today()
    if raw.lower() == "last":
        anchor = (today.replace(day=1) - relativedelta(months=1))
    elif raw == "":
        anchor = today
    else:
        try:
            dt = datetime.strptime(raw, "%Y-%m")
            anchor = dt.date()
        except ValueError:
            print("Could not parse; defaulting to current month.")
            anchor = today

    # Decide month-to-date behavior
    use_mtd = False
    if anchor.year == today.year and anchor.month == today.month:
        yn = input("Use month-to-date for the current month? [Y/n]: ").strip().lower()
        use_mtd = (yn in ("", "y", "yes"))
    else:
        yn = input("Compile the FULL month (not capped to today)? [Y/n]: ").strip().lower()
        use_mtd = not (yn in ("", "y", "yes"))  # if user says Yes full month -> not MTD

    return anchor.replace(day=1), use_mtd


def load_rows_for_window(bill_start: date, use_month_to_date: bool) -> pd.DataFrame:
    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME, engine="openpyxl")

    # Coerce to date
    if DATE_COLUMN not in df.columns:
        raise ValueError(f"Expected date column '{DATE_COLUMN}' not found in Excel. Found: {list(df.columns)}")
    df[DATE_COLUMN] = pd.to_datetime(df[DATE_COLUMN], errors="coerce").dt.date

    # Month window
    to_date_limit = date.today() if use_month_to_date else None
    window = BillingWindow.for_month(bill_start, to_date=to_date_limit)

    mask = (df[DATE_COLUMN] >= window.start) & (df[DATE_COLUMN] < window.end)
    month_df = df.loc[mask].copy()

    # Sort by visitdate then cause number
    sort_keys = [DATE_COLUMN]
    if "causeno" in month_df.columns:
        sort_keys.append("causeno")
    month_df.sort_values(sort_keys, inplace=True)

    # Build a resolved column map using aliases
    resolved_map: Dict[int, str] = {}
    for col_idx, aliases in COLUMN_MAP_BASE.items():
        found = first_existing_col(month_df, aliases)
        if not found:
            raise ValueError(f"None of the aliases {aliases} exist in the Excel header. Headers present: {list(month_df.columns)}")
        resolved_map[col_idx] = found
    month_df.attrs["resolved_map"] = resolved_map
    return month_df.reset_index(drop=True)


def chunk_list(items: List[Dict[str, Any]], size: int) -> List[List[Dict[str, Any]]]:
    return [items[i:i+size] for i in range(0, len(items), size)]


def find_target_table(doc: Document) -> int:
    # Try to locate the table whose header row contains all EXPECTED_HEADERS terms (case-insensitive contains)
    for idx, t in enumerate(doc.tables):
        if not t.rows:
            continue
        header_text = " | ".join(cell.text.strip() for cell in t.rows[0].cells)
        h_lower = header_text.lower()
        if all(h.lower() in h_lower for h in EXPECTED_HEADERS):
            return idx
    return TABLE_INDEX


def replace_placeholders(doc: Document, context: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        for key, val in context.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, val)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for key, val in context.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = run.text.replace(key, val)


def format_cell_value(excel_col_name: str, val: Any) -> str:
    if pd.isna(val):
        return ""
    # Date-like formatting if value parses as a date OR if the column looks like a date field
    if any(k in excel_col_name.lower() for k in ["date", "dated", "appointed", "submitted", "submited"]):
        try:
            if isinstance(val, date):
                return val.strftime(DATE_FORMAT)
            if isinstance(val, datetime):
                return val.date().strftime(DATE_FORMAT)
            parsed = pd.to_datetime(val, errors="coerce")
            if pd.notna(parsed):
                return parsed.date().strftime(DATE_FORMAT)
        except Exception:
            pass
    return str(val)


def fill_table(doc: Document, records: List[Dict[str, Any]], resolved_map: Dict[int, str], table_index: int) -> None:
    tables = doc.tables
    if not tables:
        raise RuntimeError("Template has no tables; expected a table with a header row and 7 data rows.")

    try:
        t = tables[table_index]
    except IndexError:
        raise RuntimeError(f"Template does not have table index {table_index}; found {len(tables)} tables.")

    required_rows = 8  # header + 7 data rows
    if len(t.rows) < required_rows:
        raise RuntimeError(f"Selected table only has {len(t.rows)} rows; need at least {required_rows} (header + 7 lines).")

    # Clear existing data area (rows 1..7)
    for r in range(1, 8):
        for c in range(len(t.columns)):
            try:
                cell = t.cell(r, c)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
            except Exception:
                pass

    for row_idx, record in enumerate(records[:7], start=1):
        for col_idx, excel_col in resolved_map.items():
            val = record.get(excel_col, "")
            txt = format_cell_value(excel_col, val)
            try:
                cell = t.cell(row_idx, col_idx)
                if cell.paragraphs:
                    cell.paragraphs[0].text = txt
                else:
                    cell.add_paragraph(txt)
            except Exception:
                pass


def build_output_filename(bill_anchor: date, form_no: int) -> str:
    month_num = bill_anchor.month
    month_name_short = bill_anchor.strftime("%b")
    year = bill_anchor.year
    return FILENAME_PATTERN.format(
        BILL_MONTH_NUM=month_num,
        BILL_MONTH_NAME_SHORT=month_name_short,
        BILL_MONTH_NAME_LONG=bill_anchor.strftime("%B"),
        BILL_YEAR=year,
        FORM_NO=form_no,
    )


def main():
    # Choose the month to compile
    bill_start, use_mtd = prompt_for_month()

    print(f"Loading Excel (read-only) from {EXCEL_PATH} …")
    df = load_rows_for_window(bill_start, use_mtd)
    resolved_map: Dict[int, str] = df.attrs["resolved_map"]

    if df.empty:
        print("No rows found for the chosen period. Nothing to do.")
        return

    # Prepare output
    ensure_output_dir(OUTPUT_DIR)

    # Records and pagination (7 per form)
    records: List[Dict[str, Any]] = df.to_dict(orient="records")
    pages = chunk_list(records, 7)

    # Common placeholders (optional if your template uses them)
    context_base = {
        "{BILL_MONTH_NUM}": str(bill_start.month),
        "{BILL_MONTH_NAME_SHORT}": bill_start.strftime("%b"),
        "{BILL_MONTH_NAME_LONG}": bill_start.strftime("%B"),
        "{BILL_YEAR}": str(bill_start.year),
    }

    for idx, page_records in enumerate(pages, start=1):
        print(f"Building form {idx} with {len(page_records)} visit(s)…")
        doc = Document(TEMPLATE_PATH)

        # Auto-locate the correct table
        table_idx = find_target_table(doc)

        # Replace placeholders that may sit in headers/footers
        context = dict(context_base)
        context["{FORM_NO}"] = str(idx)
        context["{VISIT_COUNT_ON_FORM}"] = str(len(page_records))
        replace_placeholders(doc, context)

        # Fill the table rows
        fill_table(doc, page_records, resolved_map, table_idx)

        # Save
        out_name = build_output_filename(bill_start, idx)
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        print(f"  Saved: {out_path}")

    print("Done.")


if __name__ == "__main__":
    main()

