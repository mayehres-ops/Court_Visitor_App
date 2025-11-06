from docx import Document
from pathlib import Path

# Check Payment Form
payment_template = Path(r'C:\GoogleSync\GuardianShip_App\Templates\Court_Visitor_Payment_Invoice.docx')
if payment_template.exists():
    print("=" * 70)
    print("PAYMENT FORM - Checking for content controls:")
    print("=" * 70)
    doc = Document(payment_template)

    # Check for content controls (structured document tags)
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and ('MAY' in text or 'EHRESMAN' in text or 'VENDOR' in text or 'TROON' in text or 'LAKEWAY' in text):
            print(f"Paragraph {i}: {text[:100]}")

    # Check tables for signature section
    for table_num, table in enumerate(doc.tables):
        print(f"\nTable {table_num}:")
        for row_num, row in enumerate(table.rows):
            for cell_num, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text and ('MAY' in text or 'EHRESMAN' in text or 'VENDOR' in text or 'TROON' in text):
                    print(f"  Row {row_num}, Cell {cell_num}: {text[:60]}")
else:
    print("Payment template not found")

print("\n" + "=" * 70)
print("CVR FORM - Checking for content controls:")
print("=" * 70)

# Check CVR Form
cvr_template = Path(r'C:\GoogleSync\GuardianShip_App\Templates\Court Visitor Report fillable new.docx')
if cvr_template.exists():
    doc = Document(cvr_template)

    # Look for content controls
    for element in doc.element.body.iter():
        tag = element.tag
        if 'sdt' in tag.lower():  # Structured Document Tag
            print(f"Found content control: {tag}")

    # Check first few paragraphs
    print("\nFirst 10 paragraphs with content:")
    for i, para in enumerate(doc.paragraphs[:10]):
        if para.text.strip():
            print(f"Para {i}: {para.text.strip()[:60]}")
else:
    print("CVR template not found")
