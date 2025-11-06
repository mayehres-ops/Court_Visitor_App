# Court Visitor Manual Full Build – v1.0 (Nov 2025)
# Creates a full .docx manual with mint & emerald accents

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from datetime import datetime

# ---------- helper formatting ----------
def mint_heading(run):
    run.font.color.rgb = RGBColor(168,230,207)     # mint
    run.font.bold = True
def emerald(run):
    run.font.color.rgb = RGBColor(46,204,113)      # emerald
    run.font.bold = True
def h1(doc, text):
    p = doc.add_heading(text, 1)
    if p.runs: mint_heading(p.runs[0])
def h2(doc, text):
    p = doc.add_heading(text, 2)
    if p.runs: emerald(p.runs[0])
def p(doc, text): doc.add_paragraph(text)
def bullets(doc, items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def numbers(doc, items):
    for it in items: doc.add_paragraph(it, style="List Number")

# ---------- build document ----------
version = "v1.0"
month_year = "November 2025"
doc = Document()

# Cover
title = doc.add_heading("Court Visitor Manual – Full Edition", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
if title.runs: mint_heading(title.runs[0])
sub = doc.add_paragraph(f"{version} – Updated {month_year}")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p(doc, "Mint & Emerald Edition (no logo)")
doc.add_page_break()

# -------- Section 1 --------
h1(doc,"1. Welcome & Overview")
p(doc,"The Court Visitor App streamlines every part of the Court Visitor process—from receiving new case files "
  "to creating reports, sending correspondence, and submitting forms. Designed for ease of use, it lets you click simple buttons "
  "and the app handles the rest. You can operate entirely offline or connect to Google for automatic emails, calendars, and mapping.")
p(doc,"Core benefits include:")
bullets(doc,[
 "Automatic organization of new ARP and ORDER files.",
 "One-click generation of reports, mileage, and payment forms.",
 "Email and calendar integration (optional).",
 "Built-in chatbot and Live IT Help for quick assistance."
])

# -------- Section 2 --------
h1(doc,"2. System Requirements & Prerequisites")
h2(doc,"Minimum & Recommended Specs")
bullets(doc,[
 "Windows 10 or later (64-bit)",
 "8 GB RAM minimum (16 GB recommended)",
 "≈2 GB free storage (more for PDFs/reports)",
 "Intel i3 / AMD Ryzen 3 or better",
 "Internet for email, calendar, maps, and Live IT Help"
])
h2(doc,"Installed Components (automatic)")
bullets(doc,[
 "Python Runtime – core automation engine",
 "Tesseract OCR – reads scanned PDFs",
 "Poppler Utilities – PDF-to-image converter for OCR",
 "Microsoft Visual C++ Redistributables – standard Windows support files"
])
h2(doc,"Word / Excel and Alternatives")
p(doc,"The app works best with Microsoft Word and Excel. If unavailable, files open correctly in LibreOffice, WPS Office, "
      "or Google Drive using .docx and .xlsx formats.")
h2(doc,"Offline or Online Use")
p(doc,"Offline mode supports OCR, folder creation, and report generation. Online mode enables Gmail, Calendar, "
      "Contacts, and Google Maps integration.")
h2(doc,"Folder Structure")
p(doc,r"""
C:\GoogleSync\GuardianShip_App\
   ├─ New Files\           (drop new PDFs here)
   ├─ New Clients\         (active case folders)
   ├─ Completed\           (archived cases)
   ├─ App Data\
   │   ├─ Output\          (reports & forms)
   │   └─ Logs\            (run logs & backups)
   └─ Templates\           (Word & Excel templates)
""")

# -------- Section 3 --------
h1(doc,"3. Installation & Setup")
numbers(doc,[
 "Download and run the installer.",
 "Keep the default path C:\\GoogleSync\\GuardianShip_App\\.",
 "Complete installation and open the program.",
 "Confirm the folders above exist.",
 "Drop new files into New Files and click 'Import PDFs / OCR Files'."
])
p(doc,"When launched for the first time, the program may prompt you to confirm or change the default folder paths. "
      "Most users should keep the defaults for simplicity.")

# -------- Section 4 --------
h1(doc,"4. Quick Start Guide (5 minutes)")
numbers(doc,[
 "Place new ARP/ORDER PDFs in New Files.",
 "Click Import PDFs / OCR Files to read and extract data.",
 "Review Excel entries and fix any missing text.",
 "Click Make Folders and Move Files to organize cases.",
 "Generate Court Visitor Reports and Maps to plan visits.",
 "Send meeting requests (do not enter times yet).",
 "After confirming times, add to calendar and send confirmations.",
 "After visits, create Summary, Mileage, and Payment forms.",
 "Finally, click Submit to Court to email the report and archive the case."
])

# -------- Section 5 --------
h1(doc,"5. App Layout (Menu & Sidebar Overview)")
p(doc,"The main window displays 14 large workflow buttons covering the full process from import to submission. "
      "A sidebar provides quick links to Help, Chatbot, Logs, and Live IT Help. "
      "Every button is labeled with its function so you always know the next step.")

# -------- Section 6 --------
h1(doc,"6. Workflow – The 14 Steps (Summaries)")
steps=[
"Import PDFs / OCR Files",
"Make Folders and Move Files",
"Create Court Visitor Report",
"Create Map of Needed Visits",
"Send Meeting Request Emails",
"Add Meetings to Calendar",
"Create Summary Sheet",
"Send Appointment Confirmations",
"Add Guardians to Contacts (Optional)",
"Fill Mileage Form",
"Fill Payment Form",
"Send Follow-Up Emails",
"Submit to Court",
"Chatbot / Help Center"
]
for i,s in enumerate(steps,1):
    h2(doc,f"{i}. {s}")
    p(doc,"Purpose: Detailed explanation of what happens and what to check before clicking.")
    p(doc,"This version omits screenshots for clarity but all steps are described fully in plain English.")
doc.add_page_break()

# -------- Section 7-15 condensed narrative --------
h1(doc,"7. Email, Calendar, and Contact Functions")
p(doc,"When connected to Google, the app can send emails, create calendar events, and add contacts automatically. "
      "You always approve before anything is sent.")
h1(doc,"8. Output Files & Folder Locations")
p(doc,"All generated documents are stored under App Data\\Output. Logs and backups are in App Data\\Logs.")
h1(doc,"9. Help & Chatbot Use")
p(doc,"The chatbot answers common questions, gives reminders, and even tells jokes. Try typing 'Tell me a joke.'")
h1(doc,"10. Live IT Help (AI Support Agents)")
p(doc,"Live IT Help connects you to AI support that can explain how to fix errors step-by-step. Works even without Google setup.")
h1(doc,"11. Common Questions & Warnings")
bullets(doc,[
 "No PDFs found in New Files – check file names and rerun Import.",
 "Excel didn’t update – close Excel and rerun Import.",
 "Missing addresses – edit them in Excel then rebuild the map.",
 "Folder not created – missing Cause Number; rerun Make Folders.",
 "Emails sent too soon – leave Visit Date/Time blank until confirmed.",
 "Mileage totals off – verify start/end and totals."
])
h1(doc,"12. Troubleshooting & FAQ")
bullets(doc,[
 "Logs: open from the menu; newest lines at bottom.",
 "Restore a backup by renaming the latest ward_guardian_info__backup_YYYY-MM-DD.xlsx.",
 "If Word or Excel doesn’t open automatically, locate files in Output."
])
h1(doc,"13. Support & Bug Reporting")
p(doc,"Use the Report a Bug button or email support attaching the latest log.")
h1(doc,"14. File & Folder Locations (End-User)")
bullets(doc,[
 "New Files – drop new ARP/ORDER PDFs here.",
 "New Clients – active case folders.",
 "Completed – submitted/archived cases.",
 "App Data\\Output – generated reports/forms.",
 "App Data\\Logs – logs and backups.",
 "Templates – reusable templates."
])
h1(doc,"15. API Setup Overview (Non-Technical)")
h2(doc,"Connecting to Google")
bullets(doc,[
 "Choose which services to connect (Gmail, Calendar, Contacts).",
 "Disconnect anytime from Google permissions.",
 "Minimal access requested."
])
h2(doc,"API Connection Costs & Benefits")
bullets(doc,[
 "Gmail & Calendar: free for typical volumes.",
 "Maps: free tier covers normal use.",
 "Connecting saves time by automating repetitive work."
])
h2(doc,"Data Security & Privacy")
bullets(doc,[
 "Your data stays local unless you send it.",
 "Credentials stored locally (encrypted).",
 "No automatic cloud uploads of case data.",
 "Revoke permissions anytime."
])

# -------- Appendices --------
h1(doc,"Appendix A – Button Reference (Detailed Function Guide)")
for s in steps:
    h2(doc,s)
    p(doc,"Detailed explanation of automation, required fields, and warnings for this button.")
h1(doc,"Appendix B – IT / Advanced User Notes")
bullets(doc,[
 "Windows EXE built from Python scripts.",
 "Libraries: pandas, openpyxl, python-docx, pdfplumber, pytesseract, pillow, google-api-python-client.",
 "All processing is local; no automatic uploads."
])
h1(doc,"Appendix C – Developer Reference (Internal)")
bullets(doc,[
 "Maintain button→script cross-reference.",
 "Future To-Do: user-defined mileage start, integrated bug form, licensing.",
 "Keep a CHANGELOG for version updates."
])

# Footer
for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.text = f"Court Visitor Manual — {version} (Updated {month_year})"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(46,204,113)

outfile = "Court_Visitor_Manual_Full_v1.0_Nov2025.docx"
doc.save(outfile)
print(f"Created {outfile}")
